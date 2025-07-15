"""
Módulo de análisis de layout avanzado para preservar la estructura visual del PDF
en las exportaciones a diferentes formatos.
"""

from typing import Dict, List, Any, Tuple, Optional
import logging
from pathlib import Path
import numpy as np
from collections import defaultdict
import json

# Hacer sklearn opcional
try:
    from sklearn.cluster import KMeans
    SKLEARN_AVAILABLE = True
except ImportError:
    SKLEARN_AVAILABLE = False
    KMeans = None

logger = logging.getLogger(__name__)

class LayoutAnalyzer:
    """Analizador de layout para preservar estructura visual en exports."""
    
    def __init__(self):
        self.min_font_size_header = 12
        self.column_threshold = 50  # Pixeles para detectar columnas
        self.line_height_threshold = 20  # Pixeles para detectar espaciado
        
    def analyze_document_layout(self, content: Dict[str, Any]) -> Dict[str, Any]:
        """
        Analiza la estructura completa del documento.
        
        Args:
            content: Contenido del documento con elementos y metadatos
            
        Returns:
            Estructura de layout analizada
        """
        try:
            structure = {
                "pages": {},
                "global_structure": {
                    "has_columns": False,
                    "column_count": 1,
                    "main_font_size": 11,
                    "header_levels": []
                },
                "tables": [],
                "reading_order": []
            }
            
            elements = content.get("elements", [])
            if not elements:
                return structure
            
            # Agrupa elementos por página
            pages = self._group_by_page(elements)
            
            # Analiza cada página
            for page_num, page_elements in pages.items():
                page_analysis = self._analyze_page_layout(page_elements)
                structure["pages"][page_num] = page_analysis
            
            # Análisis global
            structure["global_structure"] = self._analyze_global_structure(structure["pages"])
            
            # Determina orden de lectura
            structure["reading_order"] = self._determine_reading_order(structure["pages"])
            
            return structure
            
        except Exception as e:
            logger.error(f"Error analyzing document layout: {str(e)}")
            return {"pages": {}, "global_structure": {}, "tables": [], "reading_order": []}
    
    def _group_by_page(self, elements: List[Dict]) -> Dict[int, List[Dict]]:
        """Agrupa elementos por número de página."""
        pages = defaultdict(list)
        for element in elements:
            page_num = element.get("page", 1)
            pages[page_num].append(element)
        return dict(pages)
    
    def _analyze_page_layout(self, elements: List[Dict]) -> Dict[str, Any]:
        """Analiza el layout de una página específica."""
        analysis = {
            "columns": [],
            "headers": [],
            "text_blocks": [],
            "tables": [],
            "spatial_relationships": [],
            "bounding_box": {"x0": 0, "y0": 0, "x1": 0, "y1": 0}
        }
        
        if not elements:
            return analysis
        
        # Extrae coordenadas y texto
        coordinates = self._extract_coordinates(elements)
        
        # Detecta columnas
        analysis["columns"] = self._detect_columns(coordinates)
        
        # Detecta headers por tamaño de fuente
        analysis["headers"] = self._detect_headers(elements)
        
        # Detecta tablas por alineación
        analysis["tables"] = self._detect_tables_by_alignment(elements)
        
        # Clasifica bloques de texto
        analysis["text_blocks"] = self._classify_text_blocks(elements)
        
        # Calcula bounding box de la página
        analysis["bounding_box"] = self._calculate_page_bounds(coordinates)
        
        # Analiza relaciones espaciales
        analysis["spatial_relationships"] = self._analyze_spatial_relationships(elements)
        
        return analysis
    
    def _extract_coordinates(self, elements: List[Dict]) -> List[Dict]:
        """Extrae coordenadas de elementos."""
        coordinates = []
        for element in elements:
            bbox = element.get("bbox", {})
            if bbox:
                coordinates.append({
                    "x0": bbox.get("x0", 0),
                    "y0": bbox.get("y0", 0),
                    "x1": bbox.get("x1", 0),
                    "y1": bbox.get("y1", 0),
                    "text": element.get("text", ""),
                    "font_size": element.get("font_size", 11),
                    "element": element
                })
        return coordinates
    
    def _detect_columns(self, coordinates: List[Dict]) -> List[Dict]:
        """Detecta columnas usando clustering de posiciones X."""
        if not coordinates:
            return []
        
        try:
            # Extrae posiciones X de inicio
            x_positions = np.array([[coord["x0"]] for coord in coordinates])
            
            # Usa clustering para detectar columnas si sklearn está disponible
            if SKLEARN_AVAILABLE:
                n_clusters = min(3, len(set(coord["x0"] for coord in coordinates)))
                if n_clusters < 2:
                    return [{"x_start": 0, "x_end": 600, "elements": coordinates}]
                
                kmeans = KMeans(n_clusters=n_clusters, random_state=42)
                clusters = kmeans.fit_predict(x_positions)
                
                # Organiza por columnas
                columns = []
                for i in range(n_clusters):
                    cluster_coords = [coord for j, coord in enumerate(coordinates) if clusters[j] == i]
                    if cluster_coords:
                        x_start = min(coord["x0"] for coord in cluster_coords)
                        x_end = max(coord["x1"] for coord in cluster_coords)
                        columns.append({
                            "x_start": x_start,
                            "x_end": x_end,
                            "elements": cluster_coords
                        })
                
                # Ordena columnas por posición X
                columns.sort(key=lambda x: x["x_start"])
                return columns
            else:
                # Fallback simple sin sklearn
                logger.warning("sklearn not available, using simple column detection")
                return [{"x_start": 0, "x_end": 600, "elements": coordinates}]
                
        except Exception as e:
            logger.error(f"Error detecting columns: {str(e)}")
            return [{"x_start": 0, "x_end": 600, "elements": coordinates}]
    
    def _detect_headers(self, elements: List[Dict]) -> List[Dict]:
        """Detecta headers por tamaño de fuente y posición."""
        headers = []
        
        # Calcula tamaño de fuente promedio
        font_sizes = [elem.get("font_size", 11) for elem in elements if elem.get("font_size")]
        avg_font_size = np.mean(font_sizes) if font_sizes else 11
        
        for element in elements:
            font_size = element.get("font_size", 11)
            text = element.get("text", "").strip()
            
            # Criterios para header
            is_larger_font = font_size > avg_font_size * 1.2
            is_short_text = len(text) < 100
            is_capitalized = text.isupper() or text.istitle()
            
            if is_larger_font and is_short_text and text:
                headers.append({
                    "text": text,
                    "font_size": font_size,
                    "level": self._determine_header_level(font_size, avg_font_size),
                    "bbox": element.get("bbox", {}),
                    "element": element
                })
        
        return headers
    
    def _determine_header_level(self, font_size: float, avg_font_size: float) -> int:
        """Determina el nivel de header basado en tamaño de fuente."""
        ratio = font_size / avg_font_size
        if ratio > 2.0:
            return 1
        elif ratio > 1.5:
            return 2
        elif ratio > 1.2:
            return 3
        else:
            return 4
    
    def _detect_tables_by_alignment(self, elements: List[Dict]) -> List[Dict]:
        """Detecta tablas por alineación de texto con algoritmo mejorado."""
        tables = []
        
        # Agrupa elementos por línea Y con tolerancia más precisa
        y_groups = defaultdict(list)
        for element in elements:
            bbox = element.get("bbox", {})
            if bbox:
                y_pos = round(bbox.get("y0", 0) / 5) * 5  # Redondea a 5px para mayor precisión
                y_groups[y_pos].append(element)
        
        # Busca filas con múltiples elementos alineados
        potential_rows = []
        for y_pos, elements_in_row in y_groups.items():
            if len(elements_in_row) >= 2:  # Al menos 2 elementos para ser fila
                # Ordena por posición X
                elements_in_row.sort(key=lambda x: x.get("bbox", {}).get("x0", 0))
                
                # Verifica si los elementos están distribuidos uniformemente (indicativo de tabla)
                x_positions = [elem.get("bbox", {}).get("x0", 0) for elem in elements_in_row]
                if self._has_regular_spacing(x_positions):
                    potential_rows.append({
                        "y_pos": y_pos,
                        "elements": elements_in_row,
                        "columns": len(elements_in_row),
                        "x_positions": x_positions
                    })
        
        # Agrupa filas consecutivas en tablas con mejor detección
        if potential_rows:
            potential_rows.sort(key=lambda x: x["y_pos"])
            current_table = [potential_rows[0]]
            expected_columns = potential_rows[0]["columns"]
            
            for i in range(1, len(potential_rows)):
                current_row = potential_rows[i]
                prev_row = potential_rows[i-1]
                
                # Criterios mejorados para determinar si pertenece a la misma tabla
                y_distance = abs(current_row["y_pos"] - prev_row["y_pos"])
                column_match = abs(current_row["columns"] - expected_columns) <= 1
                x_alignment = self._check_column_alignment(current_row["x_positions"], 
                                                         current_table[0]["x_positions"])
                
                if y_distance < 30 and column_match and x_alignment:
                    current_table.append(current_row)
                    # Actualizar número esperado de columnas
                    expected_columns = max(expected_columns, current_row["columns"])
                else:
                    # Finaliza tabla actual si tiene suficientes filas
                    if len(current_table) >= 2:
                        tables.append(self._create_enhanced_table_structure(current_table))
                    current_table = [current_row]
                    expected_columns = current_row["columns"]
            
            # Añade la última tabla
            if len(current_table) >= 2:
                tables.append(self._create_enhanced_table_structure(current_table))
        
        return tables
    
    def _has_regular_spacing(self, x_positions: List[float]) -> bool:
        """Verifica si las posiciones X tienen espaciado regular (indicativo de tabla)."""
        if len(x_positions) < 2:
            return False
        
        gaps = [x_positions[i+1] - x_positions[i] for i in range(len(x_positions) - 1)]
        if not gaps:
            return False
        
        avg_gap = sum(gaps) / len(gaps)
        # Verifica si la mayoría de gaps están cerca del promedio
        regular_gaps = sum(1 for gap in gaps if abs(gap - avg_gap) < avg_gap * 0.5)
        return regular_gaps >= len(gaps) * 0.6  # 60% de los gaps deben ser regulares
    
    def _check_column_alignment(self, current_x: List[float], reference_x: List[float]) -> bool:
        """Verifica si las columnas están alineadas entre filas."""
        if not current_x or not reference_x:
            return False
        
        # Permite cierta tolerancia en la alineación
        tolerance = 20  # píxeles
        
        # Compara posiciones con la fila de referencia
        aligned_columns = 0
        for x1 in current_x:
            for x2 in reference_x:
                if abs(x1 - x2) <= tolerance:
                    aligned_columns += 1
                    break
        
        # Al menos 50% de las columnas deben estar alineadas
        return aligned_columns >= min(len(current_x), len(reference_x)) * 0.5
    
    def _create_enhanced_table_structure(self, rows: List[Dict]) -> Dict[str, Any]:
        """Crea estructura de tabla mejorada con detección de celdas combinadas."""
        table_data = []
        max_columns = max(len(row["elements"]) for row in rows)
        
        # Crear grid normalizado
        for row in rows:
            row_data = [""] * max_columns
            elements = row["elements"]
            
            for i, elem in enumerate(elements):
                if i < max_columns:
                    text = elem.get("text", "").strip()
                    # Detectar celdas que podrían ser combinadas (texto muy largo)
                    if len(text) > 50 and i < len(elements) - 1:
                        # Posible celda combinada
                        row_data[i] = text
                        # Marcar siguiente celda como parte de la combinada
                        if i + 1 < max_columns:
                            row_data[i + 1] = ""
                    else:
                        row_data[i] = text
            
            table_data.append(row_data)
        
        # Calcular bounding box de la tabla completa
        all_elements = []
        for row in rows:
            all_elements.extend(row["elements"])
        
        bbox = self._calculate_table_bbox_enhanced(all_elements)
        
        return {
            "data": table_data,
            "rows": len(table_data),
            "cols": max_columns,
            "bbox": bbox,
            "table_type": "aligned_text",
            "confidence": self._calculate_table_confidence(rows)
        }
    
    def _calculate_table_bbox_enhanced(self, elements: List[Dict]) -> Dict[str, float]:
        """Calcula bounding box mejorado de una tabla."""
        if not elements:
            return {"x0": 0, "y0": 0, "x1": 0, "y1": 0}
        
        x_coords = []
        y_coords = []
        
        for elem in elements:
            bbox = elem.get("bbox", {})
            if bbox:
                x_coords.extend([bbox.get("x0", 0), bbox.get("x1", 0)])
                y_coords.extend([bbox.get("y0", 0), bbox.get("y1", 0)])
        
        if not x_coords or not y_coords:
            return {"x0": 0, "y0": 0, "x1": 0, "y1": 0}
        
        return {
            "x0": min(x_coords),
            "y0": min(y_coords),
            "x1": max(x_coords),
            "y1": max(y_coords)
        }
    
    def _calculate_table_confidence(self, rows: List[Dict]) -> float:
        """Calcula la confianza de que la estructura detectada es una tabla."""
        if len(rows) < 2:
            return 0.3
        
        # Factores que incrementan la confianza
        confidence = 0.5  # Base
        
        # Número de filas
        if len(rows) >= 3:
            confidence += 0.1
        if len(rows) >= 5:
            confidence += 0.1
        
        # Consistencia en número de columnas
        column_counts = [row["columns"] for row in rows]
        if len(set(column_counts)) == 1:  # Todas las filas tienen mismo número de columnas
            confidence += 0.2
        
        # Alineación regular
        if all(self._check_column_alignment(row["x_positions"], rows[0]["x_positions"]) 
               for row in rows[1:]):
            confidence += 0.1
        
        return min(confidence, 0.95)
    
    def _create_table_structure(self, rows: List[Dict]) -> Dict[str, Any]:
        """Crea estructura de tabla a partir de filas detectadas."""
        table_data = []
        
        for row in rows:
            row_data = [elem.get("text", "").strip() for elem in row["elements"]]
            table_data.append(row_data)
        
        return {
            "data": table_data,
            "rows": len(table_data),
            "cols": max(len(row) for row in table_data) if table_data else 0,
            "bbox": self._calculate_table_bbox(rows)
        }
    
    def _calculate_table_bbox(self, rows: List[Dict]) -> Dict[str, float]:
        """Calcula bounding box de una tabla."""
        all_elements = []
        for row in rows:
            all_elements.extend(row["elements"])
        
        if not all_elements:
            return {"x0": 0, "y0": 0, "x1": 0, "y1": 0}
        
        x0 = min(elem.get("bbox", {}).get("x0", 0) for elem in all_elements)
        y0 = min(elem.get("bbox", {}).get("y0", 0) for elem in all_elements)
        x1 = max(elem.get("bbox", {}).get("x1", 0) for elem in all_elements)
        y1 = max(elem.get("bbox", {}).get("y1", 0) for elem in all_elements)
        
        return {"x0": x0, "y0": y0, "x1": x1, "y1": y1}
    
    def _classify_text_blocks(self, elements: List[Dict]) -> List[Dict]:
        """Clasifica bloques de texto por tipo."""
        text_blocks = []
        
        for element in elements:
            text = element.get("text", "").strip()
            if not text:
                continue
            
            block_type = self._determine_block_type(element)
            
            text_blocks.append({
                "text": text,
                "type": block_type,
                "bbox": element.get("bbox", {}),
                "font_size": element.get("font_size", 11),
                "element": element
            })
        
        return text_blocks
    
    def _determine_block_type(self, element: Dict) -> str:
        """Determina el tipo de bloque de texto."""
        text = element.get("text", "").strip()
        font_size = element.get("font_size", 11)
        
        if font_size > 14:
            return "header"
        elif len(text) < 50 and (text.isupper() or text.istitle()):
            return "title"
        elif text.startswith("•") or text.startswith("-") or text.startswith("*"):
            return "list_item"
        elif len(text) > 200:
            return "paragraph"
        else:
            return "text"
    
    def _calculate_page_bounds(self, coordinates: List[Dict]) -> Dict[str, float]:
        """Calcula los límites de la página."""
        if not coordinates:
            return {"x0": 0, "y0": 0, "x1": 600, "y1": 800}
        
        x0 = min(coord["x0"] for coord in coordinates)
        y0 = min(coord["y0"] for coord in coordinates)
        x1 = max(coord["x1"] for coord in coordinates)
        y1 = max(coord["y1"] for coord in coordinates)
        
        return {"x0": x0, "y0": y0, "x1": x1, "y1": y1}
    
    def _analyze_spatial_relationships(self, elements: List[Dict]) -> List[Dict]:
        """Analiza relaciones espaciales entre elementos."""
        relationships = []
        
        for i, elem1 in enumerate(elements):
            for j, elem2 in enumerate(elements[i+1:], i+1):
                bbox1 = elem1.get("bbox", {})
                bbox2 = elem2.get("bbox", {})
                
                if bbox1 and bbox2:
                    relationship = self._determine_spatial_relationship(bbox1, bbox2)
                    if relationship:
                        relationships.append({
                            "element1_index": i,
                            "element2_index": j,
                            "relationship": relationship,
                            "distance": self._calculate_distance(bbox1, bbox2)
                        })
        
        return relationships
    
    def _determine_spatial_relationship(self, bbox1: Dict, bbox2: Dict) -> Optional[str]:
        """Determina la relación espacial entre dos elementos."""
        # Calcula centros
        center1_x = (bbox1["x0"] + bbox1["x1"]) / 2
        center1_y = (bbox1["y0"] + bbox1["y1"]) / 2
        center2_x = (bbox2["x0"] + bbox2["x1"]) / 2
        center2_y = (bbox2["y0"] + bbox2["y1"]) / 2
        
        # Tolerancia para alineación
        tolerance = 5
        
        # Verifica alineación
        if abs(center1_y - center2_y) < tolerance:
            return "horizontally_aligned"
        elif abs(center1_x - center2_x) < tolerance:
            return "vertically_aligned"
        elif center1_y < center2_y - tolerance:
            return "above"
        elif center1_y > center2_y + tolerance:
            return "below"
        elif center1_x < center2_x - tolerance:
            return "left_of"
        elif center1_x > center2_x + tolerance:
            return "right_of"
        
        return None
    
    def _calculate_distance(self, bbox1: Dict, bbox2: Dict) -> float:
        """Calcula distancia entre dos elementos."""
        center1_x = (bbox1["x0"] + bbox1["x1"]) / 2
        center1_y = (bbox1["y0"] + bbox1["y1"]) / 2
        center2_x = (bbox2["x0"] + bbox2["x1"]) / 2
        center2_y = (bbox2["y0"] + bbox2["y1"]) / 2
        
        return np.sqrt((center1_x - center2_x)**2 + (center1_y - center2_y)**2)
    
    def _analyze_global_structure(self, pages: Dict) -> Dict[str, Any]:
        """Analiza estructura global del documento."""
        structure = {
            "has_columns": False,
            "column_count": 1,
            "main_font_size": 11,
            "header_levels": [],
            "document_type": "single_column"
        }
        
        if not pages:
            return structure
        
        # Analiza columnas
        column_counts = []
        for page_data in pages.values():
            column_count = len(page_data.get("columns", []))
            column_counts.append(column_count)
        
        if column_counts:
            avg_columns = np.mean(column_counts)
            structure["has_columns"] = avg_columns > 1
            structure["column_count"] = int(np.round(avg_columns))
        
        # Analiza headers
        all_headers = []
        for page_data in pages.values():
            all_headers.extend(page_data.get("headers", []))
        
        if all_headers:
            levels = set(header.get("level", 1) for header in all_headers)
            structure["header_levels"] = sorted(list(levels))
        
        # Determina tipo de documento
        if structure["has_columns"]:
            structure["document_type"] = "multi_column"
        elif len(all_headers) > 5:
            structure["document_type"] = "structured_document"
        else:
            structure["document_type"] = "simple_document"
        
        return structure
    
    def _determine_reading_order(self, pages: Dict) -> List[Dict]:
        """Determina el orden de lectura del documento."""
        reading_order = []
        
        for page_num, page_data in pages.items():
            columns = page_data.get("columns", [])
            
            if len(columns) <= 1:
                # Documento de una columna - orden simple por Y
                elements = []
                for col in columns:
                    elements.extend(col.get("elements", []))
                elements.sort(key=lambda x: x.get("y0", 0))
                
                for i, elem in enumerate(elements):
                    reading_order.append({
                        "page": page_num,
                        "order": i,
                        "element": elem["element"],
                        "type": "sequential"
                    })
            else:
                # Documento multi-columna - orden por columnas
                order_index = 0
                for col in columns:
                    col_elements = col.get("elements", [])
                    col_elements.sort(key=lambda x: x.get("y0", 0))
                    
                    for elem in col_elements:
                        reading_order.append({
                            "page": page_num,
                            "order": order_index,
                            "element": elem["element"],
                            "type": "columnar"
                        })
                        order_index += 1
        
        return reading_order

# Funciones helper para generar exports con layout preservado

def generate_html_with_layout(layout_analysis: Dict, content: Dict) -> str:
    """Genera HTML con posicionamiento pixel-perfect preservando el layout original."""
    css = """
    <style>
    * { margin: 0; padding: 0; box-sizing: border-box; }
    body { 
        font-family: Arial, sans-serif; 
        background: white;
        overflow-x: auto;
        min-width: 800px;
    }
    .document-container {
        position: relative;
        background: white;
        box-shadow: 0 0 10px rgba(0,0,0,0.1);
        margin: 20px auto;
        max-width: none;
    }
    .page { 
        position: relative;
        background: white;
        margin: 0 auto 30px auto;
        box-shadow: 0 2px 10px rgba(0,0,0,0.1);
        overflow: hidden;
    }
    .page-content {
        position: relative;
        width: 100%;
        height: 100%;
    }
    .text-element {
        position: absolute;
        white-space: pre-wrap;
        word-wrap: break-word;
        line-height: 1.2;
        user-select: text;
    }
    .content-block {
        margin: 20px;
        padding: 20px;
        border: 1px solid #ddd;
        border-radius: 5px;
    }
    .table-element {
        border-collapse: collapse;
        width: 100%;
        margin: 10px 0;
    }
    .table-element th, .table-element td {
        border: 1px solid #ddd;
        padding: 8px;
        text-align: left;
    }
    .table-element th {
        background-color: #f8f9fa;
        font-weight: bold;
    }
    </style>
    """
    
    html_parts = ['<!DOCTYPE html><html><head><meta charset="UTF-8"><title>Document Export</title>', css, '</head><body>']
    
    pages = layout_analysis.get("pages", {})
    
    html_parts.append('<div class="document-container">')
    
    # Si no hay páginas estructuradas, usar contenido básico
    if not pages:
        html_parts.append('<div class="page" style="width: 595px; min-height: 842px;">')
        html_parts.append('<div class="content-block">')
        
        # Usar contenido del diccionario directamente
        if isinstance(content, dict):
            import html as html_module
            
            # Procesar título si existe
            if "title" in content and content["title"]:
                title_escaped = html_module.escape(content["title"])
                html_parts.append(f'<h1 style="margin-bottom: 20px; color: #333;">{title_escaped}</h1>')
            
            # Procesar elementos estructurados si existen
            if "elements" in content and content["elements"]:
                for element in content["elements"]:
                    if isinstance(element, dict) and "content" in element:
                        element_content = element["content"]
                        if isinstance(element_content, str):
                            text_escaped = html_module.escape(element_content)
                            html_parts.append(f'<div style="margin-bottom: 10px;">{text_escaped.replace(chr(10), "<br>")}</div>')
                        elif isinstance(element_content, dict):
                            # Manejar contenido estructurado de PyMuPDF
                            if "blocks" in element_content:
                                for block in element_content["blocks"]:
                                    if "lines" in block:
                                        for line in block["lines"]:
                                            if "spans" in line:
                                                line_text = ""
                                                for span in line["spans"]:
                                                    span_text = span.get("text", "")
                                                    if span_text.strip():
                                                        line_text += span_text
                                                if line_text.strip():
                                                    text_escaped = html_module.escape(line_text)
                                                    html_parts.append(f'<div style="margin-bottom: 5px;">{text_escaped}</div>')
            
            # Procesar texto plano si no hay elementos estructurados
            elif "text" in content and content["text"]:
                text_escaped = html_module.escape(content["text"])
                # Dividir por párrafos para mejor formato
                paragraphs = text_escaped.split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        html_parts.append(f'<div style="margin-bottom: 15px;">{para.replace(chr(10), "<br>")}</div>')
            
            # Procesar tablas
            if "tables" in content and content["tables"]:
                for i, table in enumerate(content["tables"]):
                    html_parts.append(f'<h3>Table {i+1}</h3>')
                    html_parts.append('<table class="table-element">')
                    
                    if isinstance(table, dict) and "data" in table:
                        table_data = table["data"]
                        if table_data and len(table_data) > 0:
                            # Primera fila como header
                            html_parts.append('<tr>')
                            for cell in table_data[0]:
                                cell_escaped = html_module.escape(str(cell))
                                html_parts.append(f'<th>{cell_escaped}</th>')
                            html_parts.append('</tr>')
                            
                            # Resto de filas
                            for row in table_data[1:]:
                                html_parts.append('<tr>')
                                for cell in row:
                                    cell_escaped = html_module.escape(str(cell))
                                    html_parts.append(f'<td>{cell_escaped}</td>')
                                html_parts.append('</tr>')
                    html_parts.append('</table>')
            
            # Si no hay contenido procesable, mostrar mensaje
            if (not content.get("text") and not content.get("elements") and 
                not content.get("tables") and not content.get("title")):
                html_parts.append('<div style="color: #666; font-style: italic;">No content available for export. The document may need to be processed again.</div>')
        else:
            html_parts.append('<div style="color: #666; font-style: italic;">No content available for export</div>')
        
        html_parts.append('</div>')
        html_parts.append('</div>')
    else:
        # Procesamiento con páginas estructuradas
        has_content = False
        
        for page_num, page_data in pages.items():
            page_bbox = page_data.get("bounding_box", {"x0": 0, "y0": 0, "x1": 595, "y1": 842})
            page_width = max(595, page_bbox["x1"] - page_bbox["x0"])
            page_height = max(842, page_bbox["y1"] - page_bbox["y0"])
            
            html_parts.append(f'<div class="page" style="width: {page_width}px; height: {page_height}px;">')
            html_parts.append('<div class="page-content">')
            
            # Procesar elementos con coordenadas
            all_elements = []
            for column in page_data.get("columns", []):
                all_elements.extend(column.get("elements", []))
            
            page_has_content = False
            for elem in all_elements:
                element_data = elem.get("element", {})
                text = element_data.get("text", "").strip()
                
                if not text:
                    continue
                
                x = max(0, elem.get("x0", 0) - page_bbox["x0"])
                y = max(0, elem.get("y0", 0) - page_bbox["y0"])
                font_size = element_data.get("font_size", 11)
                
                import html as html_module
                text_escaped = html_module.escape(text)
                
                style = f"left: {x}px; top: {y}px; font-size: {font_size}px; font-family: Arial;"
                html_parts.append(f'<div class="text-element" style="{style}">{text_escaped}</div>')
                page_has_content = True
                has_content = True
            
            # Si la página estructurada no tiene contenido, usar fallback
            if not page_has_content:
                html_parts.append('<div class="content-block">')
                
                if isinstance(content, dict):
                    import html as html_module
                    
                    # Procesar título si existe
                    if "title" in content and content["title"]:
                        title_escaped = html_module.escape(content["title"])
                        html_parts.append(f'<h1 style="margin-bottom: 20px; color: #333;">{title_escaped}</h1>')
                    
                    # Procesar texto plano
                    if "text" in content and content["text"]:
                        text_escaped = html_module.escape(content["text"])
                        paragraphs = text_escaped.split('\n\n')
                        for para in paragraphs:
                            if para.strip():
                                html_parts.append(f'<div style="margin-bottom: 15px;">{para.replace(chr(10), "<br>")}</div>')
                    
                    # Procesar tablas
                    if "tables" in content and content["tables"]:
                        for i, table in enumerate(content["tables"]):
                            html_parts.append(f'<h3>Table {i+1}</h3>')
                            html_parts.append('<table class="table-element">')
                            
                            if isinstance(table, dict) and "data" in table:
                                table_data = table["data"]
                                if table_data and len(table_data) > 0:
                                    # Primera fila como header
                                    html_parts.append('<tr>')
                                    for cell in table_data[0]:
                                        cell_escaped = html_module.escape(str(cell))
                                        html_parts.append(f'<th>{cell_escaped}</th>')
                                    html_parts.append('</tr>')
                                    
                                    # Resto de filas
                                    for row in table_data[1:]:
                                        html_parts.append('<tr>')
                                        for cell in row:
                                            cell_escaped = html_module.escape(str(cell))
                                            html_parts.append(f'<td>{cell_escaped}</td>')
                                        html_parts.append('</tr>')
                            html_parts.append('</table>')
                
                html_parts.append('</div>')
                has_content = True
            
            html_parts.append('</div>')
            html_parts.append('</div>')
        
        # Si ninguna página tuvo contenido, mostrar mensaje
        if not has_content:
            html_parts.append('<div class="page" style="width: 595px; min-height: 842px;">')
            html_parts.append('<div class="content-block">')
            html_parts.append('<div style="color: #666; font-style: italic;">No content could be extracted from the document structure. Please try uploading the document again.</div>')
            html_parts.append('</div>')
            html_parts.append('</div>')
    
    html_parts.append('</div>')
    html_parts.append("</body></html>")
    
    return '\n'.join(html_parts)

def generate_text_with_layout(layout_analysis: Dict, content: Dict) -> str:
    """Genera texto preservando el layout original."""
    text_parts = []
    
    pages = layout_analysis.get("pages", {})
    reading_order = layout_analysis.get("reading_order", [])
    
    current_page = None
    
    for order_item in reading_order:
        page_num = order_item.get("page", 1)
        element = order_item.get("element", {})
        
        # Añade header de página
        if current_page != page_num:
            current_page = page_num
            text_parts.append(f"\n{'='*60}")
            text_parts.append(f"PAGE {page_num}")
            text_parts.append(f"{'='*60}\n")
        
        # Añade texto del elemento
        text = element.get("text", "").strip()
        if text:
            # Calcula indentación basada en posición X
            bbox = element.get("bbox", {})
            x_pos = bbox.get("x0", 0)
            indent = max(0, int(x_pos / 20))  # Aproximadamente 1 espacio por 20px
            
            # Añade espaciado vertical si es necesario
            if "header" in element.get("type", ""):
                text_parts.append(f"\n{' ' * indent}{text.upper()}")
                text_parts.append(f"{' ' * indent}{'-' * len(text)}")
            else:
                text_parts.append(f"{' ' * indent}{text}")
    
    return '\n'.join(text_parts)

def generate_excel_with_layout(layout_analysis: Dict, content: Dict, file_path: str) -> str:
    """Genera Excel con posicionamiento preciso preservando el layout original."""
    import pandas as pd
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
    from openpyxl.utils.dataframe import dataframe_to_rows
    
    # Crear workbook
    wb = Workbook()
    
    pages = layout_analysis.get("pages", {})
    
    for page_num, page_data in pages.items():
        # Crear hoja para cada página
        if page_num == 1:
            ws = wb.active
            ws.title = f"Page_{page_num}"
        else:
            ws = wb.create_sheet(title=f"Page_{page_num}")
        
        # Configurar dimensiones de página
        page_bbox = page_data.get("bounding_box", {"x0": 0, "y0": 0, "x1": 595, "y1": 842})
        page_width = page_bbox["x1"] - page_bbox["x0"]
        page_height = page_bbox["y1"] - page_bbox["y0"]
        
        # Crear grid virtual basado en coordenadas
        CELL_WIDTH = 20  # Pixeles por celda aproximadamente
        CELL_HEIGHT = 15  # Altura por fila
        
        max_col = max(50, int(page_width / CELL_WIDTH) + 5)
        max_row = max(50, int(page_height / CELL_HEIGHT) + 5)
        
        # Configurar anchos de columna
        for col in range(1, max_col + 1):
            ws.column_dimensions[ws.cell(row=1, column=col).column_letter].width = 3
        
        # Configurar alturas de fila
        for row in range(1, max_row + 1):
            ws.row_dimensions[row].height = 12
        
        # Procesar elementos de texto
        all_elements = []
        for column in page_data.get("columns", []):
            all_elements.extend(column.get("elements", []))
        
        all_elements.sort(key=lambda x: (x.get("y0", 0), x.get("x0", 0)))
        
        for elem in all_elements:
            element_data = elem.get("element", {})
            text = element_data.get("text", "").strip()
            
            if not text:
                continue
            
            # Convertir coordenadas a posición de celda
            x = elem.get("x0", 0) - page_bbox["x0"]
            y = elem.get("y0", 0) - page_bbox["y0"]
            
            col_pos = max(1, int(x / CELL_WIDTH) + 1)
            row_pos = max(1, int(y / CELL_HEIGHT) + 1)
            
            # Asegurar que está dentro de los límites
            col_pos = min(col_pos, max_col)
            row_pos = min(row_pos, max_row)
            
            # Encontrar celda vacía más cercana
            while row_pos <= max_row and ws.cell(row=row_pos, column=col_pos).value:
                row_pos += 1
            
            if row_pos <= max_row:
                cell = ws.cell(row=row_pos, column=col_pos)
                cell.value = text
                
                # Aplicar formato
                font_size = element_data.get("font_size", 11)
                font_weight = "bold" if element_data.get("flags", 0) & 16 else "normal"
                is_header = font_size > 12 and len(text) < 100
                
                cell.font = Font(
                    size=max(8, min(18, font_size)),
                    bold=(font_weight == "bold" or is_header),
                    name="Arial"
                )
                
                cell.alignment = Alignment(
                    horizontal='left',
                    vertical='top',
                    wrap_text=True
                )
                
                if is_header:
                    cell.fill = PatternFill(start_color="E6E6FA", end_color="E6E6FA", fill_type="solid")
        
        # Procesar tablas
        for table in page_data.get("tables", []):
            table_bbox = table.get("bbox", {})
            if not table_bbox:
                continue
            
            x = table_bbox.get("x0", 0) - page_bbox["x0"]
            y = table_bbox.get("y0", 0) - page_bbox["y0"]
            
            start_col = max(1, int(x / CELL_WIDTH) + 1)
            start_row = max(1, int(y / CELL_HEIGHT) + 1)
            
            table_data = table.get("data", [])
            if table_data:
                for row_idx, row in enumerate(table_data):
                    for col_idx, cell_value in enumerate(row):
                        excel_row = start_row + row_idx
                        excel_col = start_col + col_idx
                        
                        if excel_row <= max_row and excel_col <= max_col:
                            cell = ws.cell(row=excel_row, column=excel_col)
                            cell.value = str(cell_value)
                            
                            # Formato de tabla
                            thin_border = Border(
                                left=Side(style='thin'),
                                right=Side(style='thin'),
                                top=Side(style='thin'),
                                bottom=Side(style='thin')
                            )
                            cell.border = thin_border
                            
                            # Header de tabla
                            if row_idx == 0:
                                cell.font = Font(bold=True)
                                cell.fill = PatternFill(start_color="D3D3D3", end_color="D3D3D3", fill_type="solid")
        
        # Añadir información de página en la primera fila
        info_cell = ws.cell(row=1, column=1)
        info_cell.value = f"Page {page_num} - Layout preserved export"
        info_cell.font = Font(bold=True, color="666666")
    
    # Guardar archivo
    wb.save(file_path)
    return file_path

def generate_docx_with_layout(layout_analysis: Dict, content: Dict, file_path: str) -> str:
    """Genera DOCX con posicionamiento absoluto preservando el layout original."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.section import WD_SECTION
    except ImportError:
        raise ImportError("python-docx is required for DOCX export. Install with: pip install python-docx")
    
    doc = Document()
    
    pages = layout_analysis.get("pages", {})
    
    for page_num, page_data in pages.items():
        # Añadir nueva sección para cada página
        if page_num > 1:
            section = doc.add_section(WD_SECTION.NEW_PAGE)
        else:
            section = doc.sections[0]
        
        # Configurar márgenes mínimos para máximo espacio
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
        # Añadir título de página
        title = doc.add_paragraph()
        title_run = title.add_run(f"Page {page_num}")
        title_run.font.size = Pt(14)
        title_run.font.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Obtener dimensiones de página
        page_bbox = page_data.get("bounding_box", {"x0": 0, "y0": 0, "x1": 595, "y1": 842})
        
        # Procesar elementos de texto agrupados por posición Y aproximada
        all_elements = []
        for column in page_data.get("columns", []):
            all_elements.extend(column.get("elements", []))
        
        # Agrupar elementos por líneas aproximadas
        lines = {}
        for elem in all_elements:
            element_data = elem.get("element", {})
            text = element_data.get("text", "").strip()
            
            if not text:
                continue
                
            y_pos = int(elem.get("y0", 0) / 20) * 20  # Agrupar por líneas de 20px
            x_pos = elem.get("x0", 0)
            
            if y_pos not in lines:
                lines[y_pos] = []
            
            lines[y_pos].append({
                "text": text,
                "x": x_pos,
                "font_size": element_data.get("font_size", 11),
                "flags": element_data.get("flags", 0)
            })
        
        # Procesar líneas ordenadas
        for y_pos in sorted(lines.keys()):
            line_elements = sorted(lines[y_pos], key=lambda x: x["x"])
            
            # Crear párrafo para la línea
            para = doc.add_paragraph()
            
            # Determinar si usar columnas (elementos muy separados en X)
            if len(line_elements) > 1:
                x_gaps = []
                for i in range(1, len(line_elements)):
                    gap = line_elements[i]["x"] - line_elements[i-1]["x"]
                    x_gaps.append(gap)
                
                # Si hay gaps grandes, usar tabulaciones
                avg_gap = sum(x_gaps) / len(x_gaps) if x_gaps else 0
                large_gaps = [gap for gap in x_gaps if gap > avg_gap * 2]
                
                if large_gaps:
                    # Elementos en columnas - usar espaciado
                    for i, elem in enumerate(line_elements):
                        if i > 0:
                            para.add_run("\t" * max(1, int(large_gaps[0] / 100)))
                        
                        run = para.add_run(elem["text"])
                        run.font.size = Pt(max(8, min(18, elem["font_size"])))
                        run.font.bold = bool(elem["flags"] & 16)
                        run.font.italic = bool(elem["flags"] & 64)
                else:
                    # Elementos en la misma línea - concatenar con espacios
                    for i, elem in enumerate(line_elements):
                        if i > 0:
                            para.add_run(" ")
                        
                        run = para.add_run(elem["text"])
                        run.font.size = Pt(max(8, min(18, elem["font_size"])))
                        run.font.bold = bool(elem["flags"] & 16)
                        run.font.italic = bool(elem["flags"] & 64)
            else:
                # Un solo elemento
                elem = line_elements[0]
                run = para.add_run(elem["text"])
                run.font.size = Pt(max(8, min(18, elem["font_size"])))
                run.font.bold = bool(elem["flags"] & 16)
                run.font.italic = bool(elem["flags"] & 64)
                
                # Si es un header grande, centrarlo
                if elem["font_size"] > 14 and len(elem["text"]) < 100:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Procesar tablas
        for table_data_info in page_data.get("tables", []):
            table_data = table_data_info.get("data", [])
            if not table_data:
                continue
            
            # Añadir espacio antes de la tabla
            doc.add_paragraph()
            
            # Crear tabla en Word
            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]) if table_data else 1)
            table.style = 'Table Grid'
            
            for row_idx, row in enumerate(table_data):
                for col_idx, cell_value in enumerate(row):
                    if col_idx < len(table.rows[row_idx].cells):
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = str(cell_value)
                        
                        # Formato de header
                        if row_idx == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
    
    # Guardar documento
    doc.save(file_path)
    return file_path
