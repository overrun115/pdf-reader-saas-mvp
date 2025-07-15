from fastapi import APIRouter, HTTPException, UploadFile, File, Request, BackgroundTasks
from fastapi.responses import JSONResponse
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
from pathlib import Path
import uuid
import json
import time
import tempfile
import os
import logging
from datetime import datetime

from app.core.database import get_db
from app.services.file_service import process_pdf_file
from app.services.concurrency_manager import pdf_concurrency_manager, TaskStatus
from app.models.schemas import FileProcessResult, TableData

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/pdf-viewer", tags=["pdf-viewer"])

class PDFAnalysisResponse(BaseModel):
    file_id: str
    filename: str
    total_pages: int
    total_tables_found: int
    tables: List[Dict[str, Any]]
    text_blocks: List[Dict[str, Any]]
    processing_time: float

class PDFUploadResponse(BaseModel):
    file_id: str
    filename: str
    task_id: str
    status: str
    message: str
    queue_position: int

class TaskStatusResponse(BaseModel):
    task_id: str
    status: str
    filename: str
    created_at: str
    started_at: Optional[str] = None
    completed_at: Optional[str] = None
    progress_message: str
    error: Optional[str] = None
    result: Optional[PDFAnalysisResponse] = None

class QueueStatusResponse(BaseModel):
    queue_length: int
    running_tasks: int
    max_concurrent: int
    max_queue_size: int
    your_tasks: List[str]

class ExtractionRequest(BaseModel):
    file_id: str
    selected_tables: List[int]
    selected_text_blocks: Optional[List[int]] = []
    include_headers: bool = True
    output_format: str = "excel"

@router.post("/upload", response_model=PDFUploadResponse)
async def upload_pdf_for_analysis(
    file: UploadFile = File(...),
    request: Request = None
) -> PDFUploadResponse:
    """Upload PDF and queue for analysis."""
    
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Only PDF files are allowed")
    
    if file.size > 50 * 1024 * 1024:  # 50MB limit
        raise HTTPException(status_code=400, detail="File size must be less than 50MB")
    
    file_id = str(uuid.uuid4())
    
    try:
        # Save uploaded file temporarily
        temp_dir = Path(tempfile.gettempdir()) / "pdf_analysis"
        temp_dir.mkdir(exist_ok=True)
        
        file_path = temp_dir / f"{file_id}.pdf"
        with open(file_path, "wb") as f:
            content = await file.read()
            f.write(content)
        
        # Store file metadata
        cache_file = temp_dir / f"{file_id}_metadata.json"
        with open(cache_file, "w") as f:
            json.dump({
                "file_path": str(file_path),
                "filename": file.filename,
                "file_id": file_id,
                "uploaded_at": datetime.now().isoformat()
            }, f)
        
        # Submit to processing queue
        task_id = await pdf_concurrency_manager.submit_task(
            file_id=file_id,
            filename=file.filename,
            processor_func=process_pdf_with_queue,
            timeout_seconds=600  # 10 minutes timeout
        )
        
        # Get queue status
        queue_status = await pdf_concurrency_manager.get_queue_status()
        
        return PDFUploadResponse(
            file_id=file_id,
            filename=file.filename,
            task_id=task_id,
            status="queued",
            message=f"PDF uploaded successfully. Processing queued.",
            queue_position=queue_status["queue_length"]
        )
        
    except Exception as e:
        logger.error(f"Error uploading PDF: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to upload PDF: {str(e)}")

async def process_pdf_with_queue(file_id: str, filename: str) -> PDFAnalysisResponse:
    """Process PDF file through the queue system."""
    
    try:
        # Load file metadata
        temp_dir = Path(tempfile.gettempdir()) / "pdf_analysis"
        metadata_file = temp_dir / f"{file_id}_metadata.json"
        
        if not metadata_file.exists():
            raise Exception(f"Metadata file not found for {file_id}")
        
        with open(metadata_file, "r") as f:
            metadata = json.load(f)
        
        file_path = metadata["file_path"]
        
        # Enhanced PDF analysis
        analysis_result = await analyze_pdf_structure(file_path, filename)
        
        # Store analysis result
        analysis_file = temp_dir / f"{file_id}_analysis.json"
        with open(analysis_file, "w") as f:
            json.dump({
                **metadata,
                "analysis": analysis_result,
                "processed_at": datetime.now().isoformat()
            }, f)
        
        return PDFAnalysisResponse(
            file_id=file_id,
            filename=filename,
            total_pages=analysis_result.get("total_pages", 1),
            total_tables_found=len(analysis_result.get("tables", [])),
            tables=analysis_result.get("tables", []),
            text_blocks=analysis_result.get("text_blocks", []),
            processing_time=0  # Will be calculated by the manager
        )
        
    except Exception as e:
        logger.error(f"Error processing PDF {file_id}: {str(e)}")
        raise

@router.get("/task/{task_id}", response_model=TaskStatusResponse)
async def get_task_status(task_id: str) -> TaskStatusResponse:
    """Get the status of a processing task."""
    
    task = await pdf_concurrency_manager.get_task_status(task_id)
    
    if not task:
        raise HTTPException(status_code=404, detail="Task not found")
    
    # Generate progress message
    if task.status == TaskStatus.PENDING:
        progress_message = "Task is waiting in queue..."
    elif task.status == TaskStatus.RUNNING:
        progress_message = "Processing PDF structure and extracting tables..."
    elif task.status == TaskStatus.COMPLETED:
        progress_message = "Analysis completed successfully!"
    elif task.status == TaskStatus.FAILED:
        progress_message = f"Analysis failed: {task.error}"
    elif task.status == TaskStatus.TIMEOUT:
        progress_message = "Analysis timed out. Please try with a smaller file."
    else:
        progress_message = "Unknown status"
    
    return TaskStatusResponse(
        task_id=task_id,
        status=task.status.value,
        filename=task.filename,
        created_at=task.created_at.isoformat(),
        started_at=task.started_at.isoformat() if task.started_at else None,
        completed_at=task.completed_at.isoformat() if task.completed_at else None,
        progress_message=progress_message,
        error=task.error,
        result=task.result if task.status == TaskStatus.COMPLETED else None
    )

@router.get("/queue/status", response_model=QueueStatusResponse)
async def get_queue_status() -> QueueStatusResponse:
    """Get current queue status."""
    
    status = await pdf_concurrency_manager.get_queue_status()
    
    return QueueStatusResponse(
        queue_length=status["queue_length"],
        running_tasks=status["running_tasks"],
        max_concurrent=status["max_concurrent"],
        max_queue_size=status["max_queue_size"],
        your_tasks=status["pending_tasks"] + status["running_task_ids"]
    )

async def analyze_pdf_structure(file_path: str, filename: str) -> Dict[str, Any]:
    """Enhanced PDF analysis that returns detailed structure."""
    
    try:
        # Import PDF processing functions
        import sys
        SHARED_PATH = os.path.abspath(os.path.join(os.path.dirname(__file__), '../../../shared'))
        if SHARED_PATH not in sys.path:
            sys.path.insert(0, SHARED_PATH)
        
        from docling.document_converter import DocumentConverter
        from pdf_extractor import get_table_preview
        
        # Get detailed table analysis
        table_preview = get_table_preview(file_path, max_rows=50)  # More rows for better analysis
        
        # Get document structure using Docling
        converter = DocumentConverter()
        result = converter.convert(file_path)
        doc_content = result.document
        
        # Extract detailed information
        analysis = {
            "total_pages": getattr(doc_content, 'page_count', 1),
            "tables": [],
            "text_blocks": []
        }
        
        # Process tables with enhanced information
        if table_preview and table_preview.get("tables"):
            for i, table_info in enumerate(table_preview["tables"]):
                table_data = {
                    "table_id": i,
                    "page": table_info.get("page", 1),
                    "confidence": table_info.get("confidence", 0.8),
                    "columns": table_info.get("columns", []),
                    "sample_data": table_info.get("sample_data", [])[:10],  # First 10 rows for preview
                    "total_rows": len(table_info.get("sample_data", [])),
                    "bbox": table_info.get("bbox", {}),  # Bounding box coordinates
                    "table_type": detect_table_type(table_info.get("sample_data", [])),
                    "suggested_name": suggest_table_name(table_info.get("columns", []), table_info.get("sample_data", []))
                }
                analysis["tables"].append(table_data)
        
        # Extract text blocks using Docling
        if hasattr(doc_content, 'texts') or hasattr(doc_content, 'export_to_text'):
            try:
                if hasattr(doc_content, 'export_to_text'):
                    full_text = doc_content.export_to_text()
                    # Split into logical blocks
                    text_blocks = extract_text_blocks(full_text)
                    analysis["text_blocks"] = text_blocks
            except Exception as e:
                logger.warning(f"Could not extract text blocks: {e}")
                analysis["text_blocks"] = []
        
        return analysis
        
    except Exception as e:
        logger.error(f"Error in PDF structure analysis: {str(e)}")
        return {
            "total_pages": 1,
            "tables": [],
            "text_blocks": []
        }

def detect_table_type(sample_data: List[List[str]]) -> str:
    """Detect the type of table based on content."""
    if not sample_data or len(sample_data) < 2:
        return "unknown"
    
    # Check for financial data
    financial_keywords = ["amount", "price", "cost", "total", "subtotal", "tax", "$", "€", "£"]
    if any(any(keyword.lower() in str(cell).lower() for keyword in financial_keywords) 
           for row in sample_data[:3] for cell in row):
        return "financial"
    
    # Check for date-based data
    date_keywords = ["date", "time", "year", "month", "day"]
    if any(any(keyword.lower() in str(cell).lower() for keyword in date_keywords) 
           for row in sample_data[:2] for cell in row):
        return "temporal"
    
    # Check for inventory/product data
    product_keywords = ["product", "item", "sku", "inventory", "stock", "quantity"]
    if any(any(keyword.lower() in str(cell).lower() for keyword in product_keywords) 
           for row in sample_data[:2] for cell in row):
        return "inventory"
    
    return "general"

def suggest_table_name(columns: List[str], sample_data: List[List[str]]) -> str:
    """Suggest a meaningful name for the table."""
    if not columns and not sample_data:
        return "Table"
    
    # Use first column name if meaningful
    if columns and len(columns) > 0:
        first_col = columns[0].lower()
        if any(keyword in first_col for keyword in ["product", "item", "customer", "employee", "invoice"]):
            return f"{columns[0].title()} Data"
    
    # Check content for context
    if sample_data and len(sample_data) > 1:
        first_row = [str(cell).lower() for cell in sample_data[0]]
        if any("financial" in cell or "revenue" in cell or "sales" in cell for cell in first_row):
            return "Financial Data"
        elif any("inventory" in cell or "stock" in cell for cell in first_row):
            return "Inventory Data"
        elif any("customer" in cell or "client" in cell for cell in first_row):
            return "Customer Data"
    
    return "Data Table"

def extract_text_blocks(full_text: str) -> List[Dict[str, Any]]:
    """Extract logical text blocks from full text."""
    blocks = []
    
    # Split by double newlines (paragraphs)
    paragraphs = full_text.split('\n\n')
    
    for i, paragraph in enumerate(paragraphs):
        if paragraph.strip():
            block = {
                "block_id": i,
                "text": paragraph.strip(),
                "type": classify_text_block(paragraph.strip()),
                "word_count": len(paragraph.split()),
                "preview": paragraph.strip()[:100] + "..." if len(paragraph.strip()) > 100 else paragraph.strip()
            }
            blocks.append(block)
    
    return blocks[:20]  # Limit to first 20 blocks

def classify_text_block(text: str) -> str:
    """Classify text block type."""
    text_lower = text.lower()
    
    if any(keyword in text_lower for keyword in ["title", "heading", "chapter"]):
        return "heading"
    elif any(keyword in text_lower for keyword in ["summary", "conclusion", "abstract"]):
        return "summary"
    elif len(text.split()) < 10:
        return "short_text"
    else:
        return "paragraph"

@router.post("/extract")
async def extract_selected_content(
    extraction_request: ExtractionRequest,
    background_tasks: BackgroundTasks
):
    """Extract selected tables and text blocks in specified format."""
    
    try:
        # Load analysis from cache
        temp_dir = Path(tempfile.gettempdir()) / "pdf_analysis"
        cache_file = temp_dir / f"{extraction_request.file_id}_analysis.json"
        
        if not cache_file.exists():
            raise HTTPException(status_code=404, detail="Analysis not found or expired")
        
        with open(cache_file, "r") as f:
            cached_data = json.load(f)
        
        file_path = cached_data["file_path"]
        filename = cached_data["filename"]
        analysis = cached_data["analysis"]
        
        # Extract selected content
        extracted_content = await extract_content(
            file_path, 
            filename,
            analysis,
            extraction_request.selected_tables,
            extraction_request.selected_text_blocks,
            extraction_request.include_headers,
            extraction_request.output_format
        )
        
        return extracted_content
        
    except Exception as e:
        logger.error(f"Error extracting content: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to extract content: {str(e)}")

async def extract_content(
    file_path: str,
    filename: str, 
    analysis: Dict[str, Any],
    selected_tables: List[int],
    selected_text_blocks: List[int],
    include_headers: bool,
    output_format: str
) -> Dict[str, Any]:
    """Extract and format selected content."""
    
    # Re-process with full data for selected tables
    result = await process_pdf_file(file_path, filename)
    
    # Filter to selected tables
    if selected_tables:
        filtered_tables = [result.tables[i] for i in selected_tables if i < len(result.tables)]
        result.tables = filtered_tables
    
    # Create output based on format
    if output_format == "excel":
        output_path = await create_excel_output(result.tables, filename, include_headers)
    elif output_format == "csv":
        output_path = await create_csv_output(result.tables, filename, include_headers)
    elif output_format == "word":
        output_path = await create_word_output(result.tables, selected_text_blocks, analysis, filename, include_headers)
    else:
        raise HTTPException(status_code=400, detail="Unsupported output format")
    
    return {
        "download_url": f"/api/pdf-viewer/download/{os.path.basename(output_path)}",
        "filename": os.path.basename(output_path),
        "format": output_format,
        "tables_count": len(result.tables),
        "text_blocks_count": len(selected_text_blocks)
    }

async def create_excel_output(tables: List[Any], filename: str, include_headers: bool) -> str:
    """Create Excel file with selected tables."""
    import pandas as pd
    from pathlib import Path
    import tempfile
    
    temp_dir = Path(tempfile.gettempdir()) / "pdf_extractions"
    temp_dir.mkdir(exist_ok=True)
    
    output_path = temp_dir / f"{filename.replace('.pdf', '')}_extracted.xlsx"
    
    with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
        for i, table in enumerate(tables):
            table_data = table.data if hasattr(table, 'data') else table.get('data', [])
            if table_data:
                df = pd.DataFrame(table_data)
                if not include_headers and len(df) > 0:
                    df = df.iloc[1:]  # Skip first row if it contains headers
                
                sheet_name = f"Table_{i+1}"
                df.to_excel(writer, sheet_name=sheet_name, index=False, header=include_headers)
    
    return str(output_path)

async def create_csv_output(tables: List[Any], filename: str, include_headers: bool) -> str:
    """Create CSV file with combined tables."""
    import pandas as pd
    from pathlib import Path
    import tempfile
    
    temp_dir = Path(tempfile.gettempdir()) / "pdf_extractions"
    temp_dir.mkdir(exist_ok=True)
    
    output_path = temp_dir / f"{filename.replace('.pdf', '')}_extracted.csv"
    
    all_data = []
    for i, table in enumerate(tables):
        table_data = table.data if hasattr(table, 'data') else table.get('data', [])
        if table_data:
            df = pd.DataFrame(table_data)
            if not include_headers and len(df) > 0:
                df = df.iloc[1:]
            df['Table_Source'] = f"Table_{i+1}"
            all_data.append(df)
    
    if all_data:
        combined_df = pd.concat(all_data, ignore_index=True)
        combined_df.to_csv(output_path, index=False, header=include_headers)
    
    return str(output_path)

async def create_word_output(tables: List[Any], selected_text_blocks: List[int], analysis: Dict[str, Any], filename: str, include_headers: bool) -> str:
    """Create Word document with tables and text."""
    from docx import Document
    from pathlib import Path
    import tempfile
    
    temp_dir = Path(tempfile.gettempdir()) / "pdf_extractions"
    temp_dir.mkdir(exist_ok=True)
    
    output_path = temp_dir / f"{filename.replace('.pdf', '')}_extracted.docx"
    
    doc = Document()
    doc.add_heading(f'Extracted Content from {filename}', level=1)
    
    # Add selected text blocks first
    if selected_text_blocks and analysis.get("text_blocks"):
        doc.add_heading('Text Content', level=2)
        for block_id in selected_text_blocks:
            if block_id < len(analysis["text_blocks"]):
                block = analysis["text_blocks"][block_id]
                doc.add_paragraph(block["text"])
                doc.add_paragraph()  # Space between blocks
    
    # Add tables
    if tables:
        doc.add_heading('Tables', level=2)
        for i, table in enumerate(tables):
            doc.add_heading(f'Table {i+1}', level=3)
            table_data = table.data if hasattr(table, 'data') else table.get('data', [])
            
            if table_data:
                # Create Word table
                doc_table = doc.add_table(rows=len(table_data), cols=len(table_data[0]) if table_data else 1)
                doc_table.style = 'Table Grid'
                
                start_row = 0 if include_headers else 1
                for row_idx, row_data in enumerate(table_data[start_row:], start_row):
                    for col_idx, cell_data in enumerate(row_data):
                        if row_idx < len(doc_table.rows) and col_idx < len(doc_table.rows[row_idx].cells):
                            doc_table.rows[row_idx].cells[col_idx].text = str(cell_data)
                
                doc.add_paragraph()  # Space after table
    
    doc.save(output_path)
    return str(output_path)

@router.get("/download/{filename}")
async def download_extracted_file(filename: str):
    """Download extracted file."""
    from fastapi.responses import FileResponse
    
    temp_dir = Path(tempfile.gettempdir()) / "pdf_extractions"
    file_path = temp_dir / filename
    
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")
    
    return FileResponse(
        path=file_path,
        filename=filename,
        media_type='application/octet-stream'
    )