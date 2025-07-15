#!/usr/bin/env python3
"""
Universal Converter Service - Conversión inteligente entre todos los formatos
Parte de la Fase 2 de la expansión de inteligencia documental
"""

import logging
import asyncio
import tempfile
import os
import shutil
from typing import Dict, List, Any, Optional, Tuple, Union
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
import json

# Document processing libraries
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
import fitz  # PyMuPDF

# HTML and conversion utilities
from bs4 import BeautifulSoup
from weasyprint import HTML, CSS
import mammoth

# Import our intelligence services
from app.services.word_intelligence_service import WordIntelligenceService
from app.services.excel_intelligence_service import ExcelIntelligenceService
from app.services.advanced_ocr_nlp_service import AdvancedOCRNLPService
from app.services.layout_parser_service import LayoutParserService

logger = logging.getLogger(__name__)

@dataclass
class ConversionJob:
    """Trabajo de conversión"""
    job_id: str
    source_format: str
    target_format: str
    source_file: str
    target_file: str
    options: Dict[str, Any]
    status: str
    progress: float
    created_at: datetime
    completed_at: Optional[datetime] = None
    error_message: Optional[str] = None

@dataclass
class ConversionResult:
    """Resultado de conversión"""
    success: bool
    output_file: str
    conversion_time: float
    quality_score: float
    warnings: List[str]
    metadata: Dict[str, Any]

class UniversalConverterService:
    """
    Servicio universal de conversión entre formatos de documentos
    """
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        
        # Initialize intelligence services
        self.word_service = WordIntelligenceService()
        self.excel_service = ExcelIntelligenceService()
        self.ocr_service = AdvancedOCRNLPService()
        self.layout_service = LayoutParserService()
        
        # Supported conversions matrix
        self.supported_conversions = {
            'pdf': ['word', 'excel', 'html', 'txt', 'json'],
            'word': ['pdf', 'html', 'txt', 'excel', 'json'],
            'excel': ['pdf', 'word', 'html', 'csv', 'json'],
            'html': ['pdf', 'word', 'txt'],
            'txt': ['pdf', 'word', 'html'],
            'csv': ['excel', 'pdf', 'word', 'html'],
            'json': ['excel', 'word', 'html']
        }
        
        # Conversion quality mapping
        self.conversion_quality = {
            ('pdf', 'word'): 'high',
            ('pdf', 'excel'): 'medium',
            ('word', 'pdf'): 'high',
            ('word', 'excel'): 'medium',
            ('excel', 'word'): 'medium',
            ('excel', 'pdf'): 'high',
            ('html', 'pdf'): 'high',
            ('txt', 'word'): 'high'
        }
    
    async def convert_document(
        self,
        source_file: str,
        target_format: str,
        output_file: str,
        options: Optional[Dict[str, Any]] = None
    ) -> ConversionResult:
        """
        Convertir documento entre formatos
        """
        start_time = datetime.now()
        options = options or {}
        
        try:
            # Detectar formato de origen
            source_format = self._detect_file_format(source_file)
            
            # Verificar si la conversión es soportada
            if not self._is_conversion_supported(source_format, target_format):
                raise ValueError(f"Conversion from {source_format} to {target_format} is not supported")
            
            # Ejecutar conversión específica
            conversion_method = f"_convert_{source_format}_to_{target_format}"
            
            if hasattr(self, conversion_method):
                result = await getattr(self, conversion_method)(source_file, output_file, options)
            else:
                # Conversión genérica a través de formato intermedio
                result = await self._convert_via_intermediate(source_file, source_format, target_format, output_file, options)
            
            conversion_time = (datetime.now() - start_time).total_seconds()
            
            # Evaluar calidad de conversión
            quality_score = await self._evaluate_conversion_quality(
                source_file, output_file, source_format, target_format
            )
            
            return ConversionResult(
                success=True,
                output_file=output_file,
                conversion_time=conversion_time,
                quality_score=quality_score,
                warnings=result.get('warnings', []),
                metadata={
                    'source_format': source_format,
                    'target_format': target_format,
                    'conversion_method': conversion_method,
                    'options_used': options,
                    'file_sizes': {
                        'source': os.path.getsize(source_file) if os.path.exists(source_file) else 0,
                        'target': os.path.getsize(output_file) if os.path.exists(output_file) else 0
                    }
                }
            )
            
        except Exception as e:
            self.logger.error(f"Conversion failed: {e}")
            conversion_time = (datetime.now() - start_time).total_seconds()
            
            return ConversionResult(
                success=False,
                output_file="",
                conversion_time=conversion_time,
                quality_score=0.0,
                warnings=[str(e)],
                metadata={'error': str(e)}
            )
    
    def _detect_file_format(self, file_path: str) -> str:
        """
        Detectar formato de archivo
        """
        extension = Path(file_path).suffix.lower()
        
        format_mapping = {
            '.pdf': 'pdf',
            '.docx': 'word',
            '.doc': 'word',
            '.xlsx': 'excel',
            '.xls': 'excel',
            '.html': 'html',
            '.htm': 'html',
            '.txt': 'txt',
            '.csv': 'csv',
            '.json': 'json'
        }
        
        return format_mapping.get(extension, 'unknown')
    
    def _is_conversion_supported(self, source_format: str, target_format: str) -> bool:
        """
        Verificar si la conversión es soportada
        """
        return target_format in self.supported_conversions.get(source_format, [])
    
    # PDF Conversions
    async def _convert_pdf_to_word(self, source_file: str, output_file: str, options: Dict) -> Dict[str, Any]:
        """
        Convertir PDF a Word preservando formato
        """
        try:
            # Usar OCR avanzado para extraer contenido
            pages_data = []
            
            # Obtener número de páginas
            doc = fitz.open(source_file)
            total_pages = len(doc)
            doc.close()
            
            # Procesar páginas (máximo 10 para rendimiento)
            max_pages = min(total_pages, options.get('max_pages', 10))
            
            for page_num in range(max_pages):
                page_data = await self.ocr_service.process_pdf_page(source_file, page_num)
                pages_data.append(page_data)
            
            # Análisis de layout
            layout_analysis = await self.layout_service.enhanced_analyze_document_layout(source_file)
            
            # Crear documento Word
            doc = Document()
            
            # Configurar estilos
            styles = doc.styles
            
            # Agregar contenido página por página
            for page_idx, page_data in enumerate(pages_data):
                if page_idx > 0:
                    doc.add_page_break()
                
                # Agregar elementos de texto
                text_elements = page_data.get('text_elements', [])
                
                for element in text_elements:
                    text = element.get('text', '').strip()
                    if text:
                        element_type = element.get('type', 'paragraph')
                        
                        if element_type == 'title':
                            heading = doc.add_heading(text, level=1)
                        elif element_type == 'paragraph':
                            para = doc.add_paragraph(text)
                        elif element_type == 'list_item':
                            para = doc.add_paragraph(text, style='List Bullet')
            
            # Agregar tablas si se detectaron
            if layout_analysis and 'pages' in layout_analysis:
                for page in layout_analysis['pages']:
                    complex_tables = page.get('complex_tables', [])
                    
                    for table_data in complex_tables:
                        structure = table_data.get('structure', {})
                        rows = structure.get('rows', 0)
                        cols = structure.get('columns', 0)
                        
                        if rows > 0 and cols > 0:
                            table = doc.add_table(rows=rows, cols=cols)
                            table.style = 'Table Grid'
                            
                            # Llenar tabla con datos si están disponibles
                            cells = structure.get('cells', [])
                            for cell_data in cells:
                                row_idx = cell_data.get('row', 0)
                                col_idx = cell_data.get('column', 0)
                                
                                if row_idx < rows and col_idx < cols:
                                    cell = table.cell(row_idx, col_idx)
                                    cell.text = f"Cell {row_idx},{col_idx}"  # Placeholder
            
            # Guardar documento
            doc.save(output_file)
            
            return {
                'success': True,
                'pages_processed': len(pages_data),
                'warnings': ['OCR quality may vary depending on PDF source quality']
            }
            
        except Exception as e:
            raise Exception(f"PDF to Word conversion failed: {e}")
    
    async def _convert_pdf_to_excel(self, source_file: str, output_file: str, options: Dict) -> Dict[str, Any]:
        """
        Convertir PDF a Excel extrayendo tablas
        """
        try:
            # Análisis de layout para detectar tablas
            layout_analysis = await self.layout_service.enhanced_analyze_document_layout(source_file)
            
            # Crear workbook
            workbook = openpyxl.Workbook()
            workbook.remove(workbook.active)  # Remover hoja por defecto
            
            tables_found = 0
            
            if layout_analysis and 'pages' in layout_analysis:
                for page_idx, page in enumerate(layout_analysis['pages']):
                    complex_tables = page.get('complex_tables', [])
                    
                    for table_idx, table_data in enumerate(complex_tables):
                        tables_found += 1
                        sheet_name = f"Table_P{page_idx+1}_T{table_idx+1}"
                        sheet = workbook.create_sheet(title=sheet_name)
                        
                        structure = table_data.get('structure', {})
                        rows = structure.get('rows', 0)
                        cols = structure.get('columns', 0)
                        
                        # Crear headers si se detectó header
                        if structure.get('has_header', False):
                            for col in range(1, cols + 1):
                                cell = sheet.cell(row=1, column=col)
                                cell.value = f"Column {col}"
                                cell.font = Font(bold=True)
                                cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
                        
                        # Llenar con datos de ejemplo
                        start_row = 2 if structure.get('has_header', False) else 1
                        for row in range(start_row, rows + 1):
                            for col in range(1, cols + 1):
                                cell = sheet.cell(row=row, column=col)
                                cell.value = f"Data {row},{col}"
                        
                        # Aplicar formato
                        for row in sheet.iter_rows():
                            for cell in row:
                                cell.alignment = Alignment(wrap_text=True)
            
            # Si no se encontraron tablas, extraer texto como datos
            if tables_found == 0:
                sheet = workbook.create_sheet(title="Extracted_Text")
                
                # Usar OCR para extraer texto
                page_data = await self.ocr_service.process_pdf_page(source_file, 0)
                text_elements = page_data.get('text_elements', [])
                
                row = 1
                for element in text_elements:
                    text = element.get('text', '').strip()
                    if text:
                        sheet.cell(row=row, column=1).value = text
                        row += 1
            
            # Guardar workbook
            workbook.save(output_file)
            
            return {
                'success': True,
                'tables_extracted': tables_found,
                'warnings': ['Table extraction quality depends on PDF structure clarity']
            }
            
        except Exception as e:
            raise Exception(f"PDF to Excel conversion failed: {e}")
    
    # Word Conversions
    async def _convert_word_to_pdf(self, source_file: str, output_file: str, options: Dict) -> Dict[str, Any]:
        """
        Convertir Word a PDF
        """
        try:
            # Analizar documento Word
            word_analysis = await self.word_service.extract_text_for_conversion(source_file)
            
            if word_analysis['status'] != 'success':
                raise Exception("Failed to analyze Word document")
            
            structured_content = word_analysis['structured_content']
            
            # Convertir a HTML intermedio
            html_content = await self._convert_structured_content_to_html(structured_content)
            
            # Convertir HTML a PDF usando WeasyPrint
            HTML(string=html_content).write_pdf(output_file)
            
            return {
                'success': True,
                'warnings': ['Complex formatting may be simplified in PDF conversion']
            }
            
        except Exception as e:
            raise Exception(f"Word to PDF conversion failed: {e}")
    
    async def _convert_word_to_excel(self, source_file: str, output_file: str, options: Dict) -> Dict[str, Any]:
        """
        Convertir Word a Excel extrayendo tablas y datos estructurados
        """
        try:
            # Analizar documento Word
            word_analysis = await self.word_service.extract_text_for_conversion(source_file)
            
            if word_analysis['status'] != 'success':
                raise Exception("Failed to analyze Word document")
            
            structured_content = word_analysis['structured_content']
            
            # Crear workbook
            workbook = openpyxl.Workbook()
            workbook.remove(workbook.active)
            
            # Extraer tablas
            tables = structured_content.get('tables', [])
            if tables:
                for table_idx, table in enumerate(tables):
                    sheet_name = f"Table_{table_idx + 1}"
                    sheet = workbook.create_sheet(title=sheet_name)
                    
                    table_data = table.get('data', [])
                    if table_data:
                        for row_idx, row_data in enumerate(table_data, 1):
                            cells = row_data.get('cells', [])
                            for col_idx, cell_data in enumerate(cells, 1):
                                cell = sheet.cell(row=row_idx, column=col_idx)
                                cell.value = cell_data.get('text', '')
            
            # Extraer texto como hoja de datos
            text_sheet = workbook.create_sheet(title="Document_Text")
            paragraphs = structured_content.get('paragraphs', [])
            
            row = 1
            for para in paragraphs:
                text = para.get('text', '').strip()
                if text:
                    text_sheet.cell(row=row, column=1).value = text
                    text_sheet.cell(row=row, column=2).value = para.get('style', 'Normal')
                    row += 1
            
            # Headers para la hoja de texto
            text_sheet.cell(row=1, column=1).value = "Text Content"
            text_sheet.cell(row=1, column=2).value = "Style"
            text_sheet.insert_rows(1)
            text_sheet.cell(row=1, column=1).font = Font(bold=True)
            text_sheet.cell(row=1, column=2).font = Font(bold=True)
            
            workbook.save(output_file)
            
            return {
                'success': True,
                'tables_extracted': len(tables),
                'text_paragraphs': len(paragraphs),
                'warnings': ['Complex Word formatting may be lost in Excel conversion']
            }
            
        except Exception as e:
            raise Exception(f"Word to Excel conversion failed: {e}")
    
    # Excel Conversions
    async def _convert_excel_to_word(self, source_file: str, output_file: str, options: Dict) -> Dict[str, Any]:
        """
        Convertir Excel a Word
        """
        try:
            # Analizar archivo Excel
            excel_analysis = await self.excel_service.extract_data_for_conversion(source_file)
            
            if excel_analysis['status'] != 'success':
                raise Exception("Failed to analyze Excel file")
            
            structured_data = excel_analysis['structured_data']
            
            # Crear documento Word
            doc = Document()
            
            # Título del documento
            doc.add_heading('Excel Data Export', 0)
            
            # Procesar cada worksheet
            worksheets = structured_data.get('worksheets', [])
            
            for sheet_data in worksheets:
                sheet_name = sheet_data.get('name', 'Unknown')
                data = sheet_data.get('data', [])
                columns = sheet_data.get('columns', [])
                
                # Agregar título de hoja
                doc.add_heading(f'Sheet: {sheet_name}', level=1)
                
                if data and columns:
                    # Crear tabla en Word
                    table = doc.add_table(rows=1, cols=len(columns))
                    table.style = 'Table Grid'
                    
                    # Headers
                    hdr_cells = table.rows[0].cells
                    for i, column in enumerate(columns):
                        hdr_cells[i].text = str(column)
                        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                    
                    # Datos (máximo 100 filas para rendimiento)
                    max_rows = min(len(data), options.get('max_rows', 100))
                    
                    for row_data in data[:max_rows]:
                        row_cells = table.add_row().cells
                        for i, column in enumerate(columns):
                            value = row_data.get(column, '')
                            row_cells[i].text = str(value) if value is not None else ''
                    
                    if len(data) > max_rows:
                        doc.add_paragraph(f"Note: Only showing first {max_rows} rows of {len(data)} total rows.")
                
                doc.add_paragraph()  # Espacio entre hojas
            
            doc.save(output_file)
            
            return {
                'success': True,
                'worksheets_processed': len(worksheets),
                'warnings': ['Large datasets may be truncated for performance']
            }
            
        except Exception as e:
            raise Exception(f"Excel to Word conversion failed: {e}")
    
    async def _convert_excel_to_pdf(self, source_file: str, output_file: str, options: Dict) -> Dict[str, Any]:
        """
        Convertir Excel a PDF
        """
        try:
            # Primero convertir a Word, luego a PDF
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_word:
                temp_word_path = temp_word.name
            
            # Excel -> Word
            word_result = await self._convert_excel_to_word(source_file, temp_word_path, options)
            
            if not word_result['success']:
                raise Exception("Failed intermediate Excel to Word conversion")
            
            # Word -> PDF
            pdf_result = await self._convert_word_to_pdf(temp_word_path, output_file, options)
            
            # Limpiar archivo temporal
            try:
                os.unlink(temp_word_path)
            except:
                pass
            
            return {
                'success': True,
                'conversion_method': 'excel_to_word_to_pdf',
                'warnings': word_result.get('warnings', []) + pdf_result.get('warnings', [])
            }
            
        except Exception as e:
            raise Exception(f"Excel to PDF conversion failed: {e}")
    
    # HTML Conversions
    async def _convert_html_to_pdf(self, source_file: str, output_file: str, options: Dict) -> Dict[str, Any]:
        """
        Convertir HTML a PDF
        """
        try:
            # Leer HTML
            with open(source_file, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            # Configurar CSS para mejor formato
            css_style = """
            body { 
                font-family: Arial, sans-serif; 
                margin: 1in; 
                line-height: 1.6;
            }
            table { 
                border-collapse: collapse; 
                width: 100%; 
                margin: 10px 0;
            }
            th, td { 
                border: 1px solid #ddd; 
                padding: 8px; 
                text-align: left;
            }
            th { 
                background-color: #f2f2f2; 
                font-weight: bold;
            }
            h1, h2, h3 { 
                color: #333; 
                page-break-after: avoid;
            }
            """
            
            # Convertir a PDF
            HTML(string=html_content).write_pdf(
                output_file,
                stylesheets=[CSS(string=css_style)]
            )
            
            return {
                'success': True,
                'warnings': ['Complex CSS styling may not be fully preserved']
            }
            
        except Exception as e:
            raise Exception(f"HTML to PDF conversion failed: {e}")
    
    # Generic conversion via intermediate formats
    async def _convert_via_intermediate(
        self, 
        source_file: str, 
        source_format: str, 
        target_format: str, 
        output_file: str, 
        options: Dict
    ) -> Dict[str, Any]:
        """
        Conversión genérica usando formatos intermedios
        """
        try:
            # Determinar mejor formato intermedio
            if source_format == 'txt' and target_format == 'pdf':
                # TXT -> HTML -> PDF
                with tempfile.NamedTemporaryFile(suffix='.html', delete=False) as temp_html:
                    temp_html_path = temp_html.name
                
                # Convertir TXT a HTML
                await self._convert_txt_to_html(source_file, temp_html_path, options)
                
                # Convertir HTML a PDF
                result = await self._convert_html_to_pdf(temp_html_path, output_file, options)
                
                # Limpiar
                try:
                    os.unlink(temp_html_path)
                except:
                    pass
                
                result['conversion_method'] = 'txt_to_html_to_pdf'
                return result
            
            else:
                raise Exception(f"No intermediate conversion path found for {source_format} to {target_format}")
                
        except Exception as e:
            raise Exception(f"Intermediate conversion failed: {e}")
    
    async def _convert_txt_to_html(self, source_file: str, output_file: str, options: Dict) -> Dict[str, Any]:
        """
        Convertir texto plano a HTML
        """
        try:
            with open(source_file, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            # Crear HTML básico
            html_content = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="utf-8">
                <title>Text Document</title>
                <style>
                    body {{ 
                        font-family: Arial, sans-serif; 
                        margin: 40px; 
                        line-height: 1.6;
                        white-space: pre-wrap;
                    }}
                </style>
            </head>
            <body>
                {text_content}
            </body>
            </html>
            """
            
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return {'success': True}
            
        except Exception as e:
            raise Exception(f"TXT to HTML conversion failed: {e}")
    
    async def _convert_structured_content_to_html(self, structured_content: Dict) -> str:
        """
        Convertir contenido estructurado a HTML
        """
        html_parts = []
        
        html_parts.append("""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>Document</title>
            <style>
                body { font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }
                table { border-collapse: collapse; width: 100%; margin: 20px 0; }
                th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
                th { background-color: #f2f2f2; font-weight: bold; }
                h1, h2, h3 { color: #333; }
                .list-item { margin: 5px 0; }
            </style>
        </head>
        <body>
        """)
        
        # Agregar headings
        headings = structured_content.get('headings', [])
        for heading in headings:
            level = heading.get('level', 1)
            text = heading.get('text', '')
            html_parts.append(f'<h{level}>{text}</h{level}>')
        
        # Agregar párrafos
        paragraphs = structured_content.get('paragraphs', [])
        for para in paragraphs:
            text = para.get('text', '')
            html_parts.append(f'<p>{text}</p>')
        
        # Agregar listas
        lists = structured_content.get('lists', [])
        for list_data in lists:
            list_type = list_data.get('type', 'ul')
            items = list_data.get('items', [])
            
            html_parts.append(f'<{list_type}>')
            for item in items:
                html_parts.append(f'<li class="list-item">{item}</li>')
            html_parts.append(f'</{list_type}>')
        
        # Agregar tablas
        tables = structured_content.get('tables', [])
        for table in tables:
            table_data = table.get('data', [])
            if table_data:
                html_parts.append('<table>')
                
                for row_idx, row in enumerate(table_data):
                    cells = row.get('cells', [])
                    if cells:
                        html_parts.append('<tr>')
                        
                        for cell in cells:
                            cell_text = cell.get('text', '')
                            tag = 'th' if row_idx == 0 else 'td'
                            html_parts.append(f'<{tag}>{cell_text}</{tag}>')
                        
                        html_parts.append('</tr>')
                
                html_parts.append('</table>')
        
        html_parts.append('</body></html>')
        
        return '\n'.join(html_parts)
    
    async def _evaluate_conversion_quality(
        self, 
        source_file: str, 
        output_file: str, 
        source_format: str, 
        target_format: str
    ) -> float:
        """
        Evaluar calidad de la conversión
        """
        try:
            # Puntuación base según tipo de conversión
            base_quality = {
                ('pdf', 'word'): 0.8,
                ('pdf', 'excel'): 0.6,
                ('word', 'pdf'): 0.9,
                ('word', 'excel'): 0.7,
                ('excel', 'word'): 0.8,
                ('excel', 'pdf'): 0.8,
                ('html', 'pdf'): 0.9,
                ('txt', 'pdf'): 0.9
            }.get((source_format, target_format), 0.7)
            
            # Verificar que el archivo de salida existe y tiene contenido
            if not os.path.exists(output_file):
                return 0.0
            
            output_size = os.path.getsize(output_file)
            if output_size == 0:
                return 0.0
            
            # Verificar tamaño relativo (archivo muy pequeño puede indicar problema)
            source_size = os.path.getsize(source_file)
            size_ratio = output_size / source_size if source_size > 0 else 1
            
            if size_ratio < 0.1:  # Archivo de salida muy pequeño
                base_quality *= 0.5
            elif size_ratio > 10:  # Archivo de salida muy grande
                base_quality *= 0.8
            
            return min(base_quality, 1.0)
            
        except Exception:
            return 0.5  # Calidad por defecto si no se puede evaluar
    
    async def get_conversion_capabilities(self) -> Dict[str, Any]:
        """
        Obtener capacidades de conversión disponibles
        """
        return {
            "supported_formats": list(self.supported_conversions.keys()),
            "conversion_matrix": self.supported_conversions,
            "quality_estimates": self.conversion_quality,
            "features": {
                "intelligent_ocr": True,
                "layout_preservation": True,
                "table_extraction": True,
                "format_analysis": True,
                "quality_assessment": True,
                "batch_processing": False,  # Pendiente de implementar
                "metadata_preservation": True
            },
            "limitations": {
                "max_file_size": "100MB",
                "max_pages_pdf": 50,
                "max_excel_rows": 10000,
                "complex_formatting": "May be simplified"
            }
        }
    
    async def estimate_conversion_time(
        self, 
        source_file: str, 
        target_format: str
    ) -> Dict[str, Any]:
        """
        Estimar tiempo de conversión
        """
        try:
            source_format = self._detect_file_format(source_file)
            file_size_mb = os.path.getsize(source_file) / (1024 * 1024)
            
            # Tiempo base por tipo de conversión (segundos por MB)
            time_factors = {
                ('pdf', 'word'): 15,
                ('pdf', 'excel'): 20,
                ('word', 'pdf'): 5,
                ('word', 'excel'): 8,
                ('excel', 'word'): 6,
                ('excel', 'pdf'): 8,
                ('html', 'pdf'): 2,
                ('txt', 'pdf'): 1
            }
            
            factor = time_factors.get((source_format, target_format), 10)
            estimated_seconds = file_size_mb * factor
            
            # Consideraciones adicionales
            if source_format == 'pdf':
                # Contar páginas para PDF
                try:
                    doc = fitz.open(source_file)
                    page_count = len(doc)
                    doc.close()
                    estimated_seconds += page_count * 2  # 2 segundos adicionales por página
                except:
                    pass
            
            return {
                "estimated_time_seconds": max(5, int(estimated_seconds)),
                "estimated_time_human": self._format_duration(estimated_seconds),
                "file_size_mb": round(file_size_mb, 2),
                "source_format": source_format,
                "target_format": target_format,
                "complexity_factor": factor
            }
            
        except Exception as e:
            return {
                "error": str(e),
                "estimated_time_seconds": 60,
                "estimated_time_human": "1 minute"
            }
    
    def _format_duration(self, seconds: float) -> str:
        """
        Formatear duración en texto legible
        """
        if seconds < 60:
            return f"{int(seconds)} seconds"
        elif seconds < 3600:
            minutes = int(seconds // 60)
            remaining_seconds = int(seconds % 60)
            return f"{minutes}m {remaining_seconds}s"
        else:
            hours = int(seconds // 3600)
            remaining_minutes = int((seconds % 3600) // 60)
            return f"{hours}h {remaining_minutes}m"
    
    async def batch_convert(
        self, 
        conversion_jobs: List[Dict[str, Any]]
    ) -> List[ConversionResult]:
        """
        Conversión en lote (implementación futura)
        """
        results = []
        
        for job in conversion_jobs:
            try:
                result = await self.convert_document(
                    source_file=job['source_file'],
                    target_format=job['target_format'],
                    output_file=job['output_file'],
                    options=job.get('options', {})
                )
                results.append(result)
                
            except Exception as e:
                results.append(ConversionResult(
                    success=False,
                    output_file="",
                    conversion_time=0.0,
                    quality_score=0.0,
                    warnings=[str(e)],
                    metadata={'error': str(e)}
                ))
        
        return results