from sqlalchemy.orm import Session
from typing import List, Optional
import os
from datetime import datetime
import asyncio

from app.models.database import ProcessedFile, FileStatus, OutputFormat, User, UserTier
from app.core.config import settings

def save_uploaded_file(file_content: bytes, filename: str) -> str:
    """Save uploaded file to disk and return file path"""
    
    upload_dir = settings.UPLOAD_DIR
    if not os.path.exists(upload_dir):
        os.makedirs(upload_dir)
    
    file_path = os.path.join(upload_dir, filename)
    
    with open(file_path, "wb") as f:
        f.write(file_content)
    
    return file_path

def create_file_record(
    db: Session,
    user_id: int,
    original_filename: str,
    file_size: int,
    file_hash: str,
    file_path: str,
    file_type: str = 'pdf'
) -> ProcessedFile:
    """Create a new file record in the database"""
    
    file_record = ProcessedFile(
        user_id=user_id,
        original_filename=original_filename,
        file_size=file_size,
        file_hash=file_hash,
        input_file_path=file_path,
        status=FileStatus.UPLOADED,
        file_type=file_type
    )
    
    db.add(file_record)
    db.commit()
    db.refresh(file_record)
    
    return file_record

def get_file_by_id(db: Session, file_id: int, user_id: int) -> Optional[ProcessedFile]:
    """Get a file record by ID and user ID"""
    
    return db.query(ProcessedFile).filter(
        ProcessedFile.id == file_id,
        ProcessedFile.user_id == user_id
    ).first()

def get_user_files(
    db: Session,
    user_id: int,
    skip: int = 0,
    limit: int = 50,
    status_filter: Optional[str] = None
) -> List[ProcessedFile]:
    """Get list of user's files with optional filtering"""
    
    query = db.query(ProcessedFile).filter(ProcessedFile.user_id == user_id)
    
    if status_filter:
        try:
            status_enum = FileStatus(status_filter)
            query = query.filter(ProcessedFile.status == status_enum)
        except ValueError:
            pass  # Invalid status filter, ignore
    
    return query.order_by(ProcessedFile.created_at.desc()).offset(skip).limit(limit).all()

def update_file_status(
    db: Session,
    file_id: int,
    status: FileStatus,
    output_file_path: Optional[str] = None,
    tables_found: Optional[int] = None,
    total_rows: Optional[int] = None,
    processing_time: Optional[float] = None,
    error_message: Optional[str] = None
) -> bool:
    """Update file processing status and results"""
    
    file_record = db.query(ProcessedFile).filter(ProcessedFile.id == file_id).first()
    
    if not file_record:
        return False
    
    file_record.status = status
    
    if output_file_path:
        file_record.output_file_path = output_file_path
    
    if tables_found is not None:
        file_record.tables_found = tables_found
    
    if total_rows is not None:
        file_record.total_rows = total_rows
    
    if processing_time is not None:
        file_record.processing_time = processing_time
    
    if error_message is not None:
        file_record.error_message = error_message
    
    if status == FileStatus.COMPLETED:
        file_record.completed_at = datetime.utcnow()
        
        # Update user's processing count
        user = db.query(User).filter(User.id == file_record.user_id).first()
        if user:
            user.files_processed_this_month += 1
            user.total_files_processed += 1
            
            # Check if we need to send usage warning email
            asyncio.create_task(check_and_send_usage_warning(user))
    
    db.commit()
    return True

def delete_file_record(db: Session, file_id: int) -> bool:
    """Delete a file record from database"""
    
    file_record = db.query(ProcessedFile).filter(ProcessedFile.id == file_id).first()
    
    if not file_record:
        return False
    
    # First delete related editable table data to avoid foreign key constraint violation
    try:
        from app.models.database import EditableTableData
        db.query(EditableTableData).filter(EditableTableData.file_id == file_id).delete()
    except Exception as e:
        print(f"Warning: Error deleting editable table data for file {file_id}: {e}")
        # Continue with file deletion even if editable table deletion fails
    
    db.delete(file_record)
    db.commit()
    return True

def get_files_by_hash(db: Session, file_hash: str, user_id: int) -> List[ProcessedFile]:
    """Get files with the same hash for deduplication"""
    
    return db.query(ProcessedFile).filter(
        ProcessedFile.file_hash == file_hash,
        ProcessedFile.user_id == user_id
    ).all()

def cleanup_old_files(db: Session, days_old: int = 30) -> int:
    """Clean up old files and their records"""
    
    from datetime import datetime, timedelta
    
    cutoff_date = datetime.utcnow() - timedelta(days=days_old)
    
    old_files = db.query(ProcessedFile).filter(
        ProcessedFile.created_at < cutoff_date
    ).all()
    
    deleted_count = 0
    
    for file_record in old_files:
        try:
            # Delete physical files
            if os.path.exists(file_record.input_file_path):
                os.remove(file_record.input_file_path)
            
            if file_record.output_file_path and os.path.exists(file_record.output_file_path):
                os.remove(file_record.output_file_path)
            
            # Delete database record
            db.delete(file_record)
            deleted_count += 1
            
        except Exception as e:
            print(f"Error deleting file {file_record.id}: {e}")
            continue
    
    db.commit()
    return deleted_count

def get_processing_statistics(db: Session, user_id: Optional[int] = None) -> dict:
    """Get processing statistics for analytics"""
    
    query = db.query(ProcessedFile)
    
    if user_id:
        query = query.filter(ProcessedFile.user_id == user_id)
    
    total_files = query.count()
    completed_files = query.filter(ProcessedFile.status == FileStatus.COMPLETED).count()
    failed_files = query.filter(ProcessedFile.status == FileStatus.FAILED).count()
    processing_files = query.filter(ProcessedFile.status == FileStatus.PROCESSING).count()
    
    # Calculate average processing time for completed files
    completed_query = query.filter(
        ProcessedFile.status == FileStatus.COMPLETED,
        ProcessedFile.processing_time.isnot(None)
    )
    
    avg_processing_time = None
    total_tables = 0
    total_rows = 0
    
    if completed_query.count() > 0:
        processing_times = [f.processing_time for f in completed_query.all() if f.processing_time]
        if processing_times:
            avg_processing_time = sum(processing_times) / len(processing_times)
        
        total_tables = sum(f.tables_found or 0 for f in completed_query.all())
        total_rows = sum(f.total_rows or 0 for f in completed_query.all())
    
    return {
        "total_files": total_files,
        "completed_files": completed_files,
        "failed_files": failed_files,
        "processing_files": processing_files,
        "success_rate": completed_files / total_files if total_files > 0 else 0,
        "avg_processing_time": avg_processing_time,
        "total_tables_extracted": total_tables,
        "total_rows_extracted": total_rows
    }

async def check_and_send_usage_warning(user: User):
    """Check if user is approaching usage limit and send warning email"""
    try:
        # Get tier limits
        tier_limits = {
            UserTier.FREE: settings.FREE_TIER_LIMIT,
            UserTier.BASIC: settings.BASIC_TIER_LIMIT,
            UserTier.PRO: settings.PRO_TIER_LIMIT,
            UserTier.ENTERPRISE: settings.ENTERPRISE_TIER_LIMIT,
        }
        
        tier_limit = tier_limits.get(user.tier, settings.FREE_TIER_LIMIT)
        
        # Skip if unlimited plan
        if tier_limit == -1:
            return
        
        # Calculate usage percentage
        usage_percentage = (user.files_processed_this_month / tier_limit) * 100
        
        # Send warning at 80% and 95% usage
        should_send_warning = False
        
        if usage_percentage >= 95 and user.files_processed_this_month % 5 == 0:
            # Send warning every 5 files after 95%
            should_send_warning = True
        elif usage_percentage >= 80 and user.files_processed_this_month == int(tier_limit * 0.8):
            # Send warning exactly at 80%
            should_send_warning = True
        
        if should_send_warning:
            from app.services.email_service import email_service
            await email_service.send_usage_limit_warning_email(user, int(usage_percentage))
            
    except Exception as e:
        # Don't let email errors affect file processing
        print(f"Error sending usage warning email: {e}")
        pass

async def process_pdf_file(file_path: str, filename: str) -> 'FileProcessResult':
    """
    Process PDF file and return structured result data.
    
    Args:
        file_path: Path to the PDF file
        filename: Original filename
        
    Returns:
        FileProcessResult: Structured processing results
    """
    import time
    from app.models.schemas import FileProcessResult, TableData
    
    start_time = time.time()
    
    try:
        # Import PDF processing functions
        import sys
        import os
        SHARED_PATH = os.path.abspath(os.path.join(os.path.dirname(__file__), '../../../shared'))
        if SHARED_PATH not in sys.path:
            sys.path.insert(0, SHARED_PATH)
        
        from pdf_extractor import get_table_preview
        
        # Get table preview data
        preview_data = get_table_preview(file_path, max_rows=1000)  # Get all rows for full processing
        
        if not preview_data or not preview_data.get("tables"):
            return FileProcessResult(
                tables=[],
                total_tables_found=0,
                processing_time=time.time() - start_time
            )
        
        # Convert preview data to FileProcessResult format
        tables = []
        for table_info in preview_data["tables"]:
            # Convert sample_data (list of dicts) to list of lists
            data_rows = []
            headers = table_info.get("columns", [])
            
            # Add headers as first row if we have them
            if headers:
                data_rows.append(headers)
            
            # Add data rows
            for row_dict in table_info.get("sample_data", []):
                row_values = [str(row_dict.get(col, "")) for col in headers]
                data_rows.append(row_values)
            
            table_data = TableData(
                data=data_rows,
                headers=headers,
                table_index=table_info["table_number"]
            )
            tables.append(table_data)
        
        processing_time = time.time() - start_time
        
        return FileProcessResult(
            tables=tables,
            total_tables_found=len(tables),
            processing_time=processing_time
        )
        
    except Exception as e:
        # Return error result
        return FileProcessResult(
            tables=[],
            total_tables_found=0,
            processing_time=time.time() - start_time
        )

def create_word_document(tables: list, filename: str, output_path: str) -> str:
    """
    Create a Word document with extracted tables.
    
    Args:
        tables: List of table data from FileProcessResult
        filename: Original PDF filename
        output_path: Path where to save the Word document
        
    Returns:
        str: Path to the created Word document
    """
    from docx import Document
    from docx.shared import Inches
    from docx.enum.table import WD_TABLE_ALIGNMENT
    import os
    
    # Create a new Document
    doc = Document()
    
    # Add title
    title = doc.add_heading(f'Tables Extracted from {filename}', level=1)
    title.alignment = 1  # Center alignment
    
    # Add generation info
    doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'Total tables found: {len(tables)}')
    doc.add_paragraph()  # Empty line
    
    # Add each table
    for i, table_data in enumerate(tables):
        # Add table heading
        heading = doc.add_heading(f'Table {i + 1}', level=2)
        
        # Get table data
        data = table_data.get('data', [])
        if not data:
            doc.add_paragraph('No data available for this table.')
            continue
            
        # Create table in Word
        table = doc.add_table(rows=len(data), cols=len(data[0]) if data else 1)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Populate table
        for row_idx, row_data in enumerate(data):
            for col_idx, cell_data in enumerate(row_data):
                if row_idx < len(table.rows) and col_idx < len(table.rows[row_idx].cells):
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = str(cell_data) if cell_data else ''
                    
                    # Make header row bold
                    if row_idx == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
        
        # Add spacing after table
        doc.add_paragraph()
    
    # Add footer
    doc.add_page_break()
    footer_para = doc.add_paragraph()
    footer_para.add_run('Generated by PDF Table Extractor').italic = True
    footer_para.alignment = 1  # Center alignment
    
    # Save document
    doc.save(output_path)
    return output_path

def create_text_document(file_path: str, filename: str, output_path: str) -> str:
    """
    Extract all text from PDF and save as Word document.
    
    Args:
        file_path: Path to the PDF file
        filename: Original PDF filename
        output_path: Path where to save the Word document
        
    Returns:
        str: Path to the created Word document
    """
    from docx import Document
    import sys
    import os
    
    try:
        # Import PDF processing functions
        SHARED_PATH = os.path.abspath(os.path.join(os.path.dirname(__file__), '../../../shared'))
        if SHARED_PATH not in sys.path:
            sys.path.insert(0, SHARED_PATH)
        
        from docling.document_converter import DocumentConverter
        
        # Convert PDF to get all text
        converter = DocumentConverter()
        result = converter.convert(file_path)
        doc_content = result.document
        
        # Create Word document
        word_doc = Document()
        
        # Add title
        title = word_doc.add_heading(f'Text Extracted from {filename}', level=1)
        title.alignment = 1  # Center alignment
        
        # Add generation info
        word_doc.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        word_doc.add_paragraph()  # Empty line
        
        # Add extracted text
        if hasattr(doc_content, 'export_to_text'):
            text_content = doc_content.export_to_text()
        else:
            # Fallback: extract text from tables and other elements
            text_content = ""
            if hasattr(doc_content, 'tables'):
                for i, table in enumerate(doc_content.tables):
                    text_content += f"\n\nTable {i+1}:\n"
                    try:
                        df = table.export_to_dataframe()
                        text_content += df.to_string(index=False)
                    except:
                        text_content += "Could not extract table data"
        
        # Add text to document with proper formatting
        paragraphs = text_content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = word_doc.add_paragraph(para_text.strip())
                
        # Add footer
        word_doc.add_page_break()
        footer_para = word_doc.add_paragraph()
        footer_para.add_run('Generated by PDF Table Extractor').italic = True
        footer_para.alignment = 1  # Center alignment
        
        # Save document
        word_doc.save(output_path)
        return output_path
        
    except Exception as e:
        raise Exception(f"Failed to create text document: {str(e)}")