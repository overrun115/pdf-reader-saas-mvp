#!/usr/bin/env python3
"""
Word Intelligence Service - Procesamiento avanzado de documentos Word
Parte de la Fase 2 de la expansión de inteligencia documental
"""

import logging
import asyncio
import io
import re
import zipfile
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
import xml.etree.ElementTree as ET

# Word processing libraries
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import docx2txt
from docx2python import docx2python

# Additional analysis
import mammoth
from bs4 import BeautifulSoup

# NLP and validation
try:
    import spacy
    SPACY_AVAILABLE = True
except ImportError:
    SPACY_AVAILABLE = False

logger = logging.getLogger(__name__)

@dataclass
class WordElement:
    """Elemento extraído de documento Word"""
    element_type: str  # paragraph, table, image, header, footer, list
    content: str
    properties: Dict[str, Any]
    position: int
    metadata: Optional[Dict[str, Any]] = None

@dataclass
class WordDocumentStructure:
    """Estructura del documento Word"""
    sections: List[Dict[str, Any]]
    elements: List[WordElement]
    tables: List[Dict[str, Any]]
    images: List[Dict[str, Any]]
    styles: Dict[str, Any]
    metadata: Dict[str, Any]

class WordIntelligenceService:
    """
    Servicio de inteligencia para documentos Word
    """
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        
        # Initialize NLP if available
        self.nlp = None
        if SPACY_AVAILABLE:
            try:
                self.nlp = spacy.load("en_core_web_sm")
            except OSError:
                self.logger.warning("spaCy model not available for Word analysis")
        
        # Word-specific patterns
        self.style_patterns = {
            'heading_1': r'^(Chapter|Section|\d+\.|\w+:)\s+',
            'heading_2': r'^\d+\.\d+\s+',
            'heading_3': r'^\d+\.\d+\.\d+\s+',
            'bullet_point': r'^\s*[•·▪▫◦‣⁃]\s*',
            'numbered_list': r'^\s*\d+[.)]\s*',
            'emphasis': r'\*\*(.+?)\*\*|__(.+?)__|_(.+?)_|\*(.+?)\*'
        }
        
    async def analyze_word_document(self, file_path: str) -> Dict[str, Any]:
        """
        Análisis completo de documento Word
        """
        try:
            start_time = datetime.now()
            
            # Análisis con múltiples librerías para máxima información
            docx_analysis = await self._analyze_with_python_docx(file_path)
            content_analysis = await self._analyze_content_structure(file_path)
            style_analysis = await self._analyze_document_styles(file_path)
            table_analysis = await self._analyze_embedded_tables(file_path)
            metadata_analysis = await self._extract_document_metadata(file_path)
            
            # Análisis de calidad y formato
            quality_analysis = await self._assess_document_quality(docx_analysis, content_analysis)
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return {
                "document_type": "word",
                "analysis_results": {
                    "structure": docx_analysis,
                    "content": content_analysis,
                    "styles": style_analysis,
                    "tables": table_analysis,
                    "metadata": metadata_analysis,
                    "quality_assessment": quality_analysis
                },
                "processing_time": processing_time,
                "status": "completed",
                "capabilities_used": self._get_analysis_capabilities()
            }
            
        except Exception as e:
            self.logger.error(f"Word document analysis failed: {e}")
            return {
                "document_type": "word",
                "status": "failed",
                "error": str(e),
                "processing_time": (datetime.now() - start_time).total_seconds() if 'start_time' in locals() else 0
            }
    
    async def _analyze_with_python_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Análisis usando python-docx para estructura detallada
        """
        try:
            doc = Document(file_path)
            
            # Analizar paragraphs
            paragraphs = []
            for i, para in enumerate(doc.paragraphs):
                para_info = {
                    "index": i,
                    "text": para.text,
                    "style": para.style.name if para.style else "Normal",
                    "alignment": self._get_alignment_name(para.alignment),
                    "runs": []
                }
                
                # Analizar runs para formato detallado
                for run in para.runs:
                    run_info = {
                        "text": run.text,
                        "bold": run.bold,
                        "italic": run.italic,
                        "underline": run.underline,
                        "font_name": run.font.name,
                        "font_size": run.font.size.pt if run.font.size else None
                    }
                    para_info["runs"].append(run_info)
                
                paragraphs.append(para_info)
            
            # Analizar tablas
            tables = []
            for i, table in enumerate(doc.tables):
                table_info = await self._analyze_word_table(table, i)
                tables.append(table_info)
            
            # Analizar secciones
            sections = []
            for i, section in enumerate(doc.sections):
                section_info = {
                    "index": i,
                    "page_width": section.page_width.inches if section.page_width else None,
                    "page_height": section.page_height.inches if section.page_height else None,
                    "orientation": "landscape" if (section.page_width and section.page_height and 
                                                 section.page_width > section.page_height) else "portrait"
                }
                sections.append(section_info)
            
            return {
                "paragraphs": paragraphs,
                "tables": tables,
                "sections": sections,
                "total_paragraphs": len(paragraphs),
                "total_tables": len(tables),
                "total_sections": len(sections)
            }
            
        except Exception as e:
            self.logger.error(f"python-docx analysis failed: {e}")
            return {"error": str(e)}
    
    async def _analyze_word_table(self, table, table_index: int) -> Dict[str, Any]:
        """
        Análisis detallado de tabla en Word
        """
        try:
            rows_data = []
            
            for row_idx, row in enumerate(table.rows):
                row_data = {
                    "row_index": row_idx,
                    "cells": []
                }
                
                for cell_idx, cell in enumerate(row.cells):
                    cell_data = {
                        "cell_index": cell_idx,
                        "text": cell.text.strip(),
                        "paragraphs": len(cell.paragraphs),
                        "width": cell.width.inches if cell.width else None
                    }
                    row_data["cells"].append(cell_data)
                
                rows_data.append(row_data)
            
            # Analizar estructura de la tabla
            table_structure = self._analyze_table_structure(rows_data)
            
            return {
                "table_index": table_index,
                "rows": len(rows_data),
                "columns": len(rows_data[0]["cells"]) if rows_data else 0,
                "data": rows_data,
                "structure_analysis": table_structure,
                "table_type": self._classify_word_table(rows_data)
            }
            
        except Exception as e:
            return {
                "table_index": table_index,
                "error": str(e),
                "rows": 0,
                "columns": 0
            }
    
    def _analyze_table_structure(self, rows_data: List[Dict]) -> Dict[str, Any]:
        """
        Analizar estructura y patrones en tabla de Word
        """
        if not rows_data:
            return {"type": "empty"}
        
        # Detectar header row
        has_header = False
        if len(rows_data) > 1:
            first_row_empty_cells = sum(1 for cell in rows_data[0]["cells"] if not cell["text"])
            second_row_empty_cells = sum(1 for cell in rows_data[1]["cells"] if not cell["text"])
            
            # Si la primera fila tiene menos celdas vacías, probablemente es header
            has_header = first_row_empty_cells < second_row_empty_cells
        
        # Detectar patrones de datos
        numeric_columns = []
        text_columns = []
        
        if len(rows_data) > 1:
            num_cols = len(rows_data[0]["cells"])
            
            for col_idx in range(num_cols):
                column_values = []
                for row in rows_data[1:]:  # Skip potential header
                    if col_idx < len(row["cells"]):
                        column_values.append(row["cells"][col_idx]["text"])
                
                # Analizar tipo de columna
                numeric_count = sum(1 for val in column_values if self._is_numeric(val))
                if numeric_count / len(column_values) > 0.7:  # 70% numérico
                    numeric_columns.append(col_idx)
                else:
                    text_columns.append(col_idx)
        
        return {
            "has_header": has_header,
            "numeric_columns": numeric_columns,
            "text_columns": text_columns,
            "total_cells": sum(len(row["cells"]) for row in rows_data),
            "empty_cells": sum(1 for row in rows_data for cell in row["cells"] if not cell["text"].strip())
        }
    
    def _classify_word_table(self, rows_data: List[Dict]) -> str:
        """
        Clasificar tipo de tabla basándose en contenido
        """
        if not rows_data or len(rows_data) < 2:
            return "simple"
        
        # Analizar contenido para clasificación
        all_text = " ".join([
            cell["text"] for row in rows_data for cell in row["cells"]
        ]).lower()
        
        # Patrones para diferentes tipos de tabla
        if any(word in all_text for word in ["total", "sum", "amount", "price", "$"]):
            return "financial"
        elif any(word in all_text for word in ["name", "email", "phone", "address"]):
            return "contact_list"
        elif any(word in all_text for word in ["date", "time", "schedule", "agenda"]):
            return "schedule"
        elif len(rows_data) > 10 and len(rows_data[0]["cells"]) > 3:
            return "data_table"
        else:
            return "generic"
    
    def _is_numeric(self, text: str) -> bool:
        """Verificar si el texto es numérico"""
        try:
            # Limpiar formato común
            cleaned = re.sub(r'[$,%\s]', '', text.strip())
            float(cleaned)
            return True
        except (ValueError, AttributeError):
            return False
    
    async def _analyze_content_structure(self, file_path: str) -> Dict[str, Any]:
        """
        Análisis de estructura de contenido usando docx2txt y mammoth
        """
        try:
            # Extraer texto plano
            plain_text = docx2txt.process(file_path)
            
            # Extraer con mammoth para HTML
            with open(file_path, "rb") as docx_file:
                result = mammoth.convert_to_html(docx_file)
                html_content = result.value
                conversion_messages = result.messages
            
            # Analizar estructura HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Detectar headings
            headings = []
            for tag in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                headings.append({
                    "level": int(tag.name[1]),
                    "text": tag.get_text().strip(),
                    "tag": tag.name
                })
            
            # Detectar listas
            lists = []
            for ul in soup.find_all(['ul', 'ol']):
                list_items = [li.get_text().strip() for li in ul.find_all('li')]
                lists.append({
                    "type": ul.name,
                    "items": list_items,
                    "item_count": len(list_items)
                })
            
            # Análisis de texto con NLP si está disponible
            nlp_analysis = {}
            if self.nlp and plain_text:
                nlp_analysis = await self._analyze_text_with_nlp(plain_text)
            
            return {
                "plain_text": plain_text,
                "html_content": html_content,
                "text_length": len(plain_text),
                "word_count": len(plain_text.split()) if plain_text else 0,
                "headings": headings,
                "lists": lists,
                "nlp_analysis": nlp_analysis,
                "conversion_messages": [str(msg) for msg in conversion_messages]
            }
            
        except Exception as e:
            self.logger.error(f"Content structure analysis failed: {e}")
            return {"error": str(e)}
    
    async def _analyze_text_with_nlp(self, text: str) -> Dict[str, Any]:
        """
        Análisis NLP del texto del documento
        """
        try:
            if not self.nlp:
                return {"error": "NLP not available"}
            
            # Limitar texto para análisis (primeros 10000 caracteres)
            text_sample = text[:10000] if len(text) > 10000 else text
            doc = self.nlp(text_sample)
            
            # Extraer entidades
            entities = []
            for ent in doc.ents:
                entities.append({
                    "text": ent.text,
                    "label": ent.label_,
                    "description": spacy.explain(ent.label_) if spacy.explain(ent.label_) else ent.label_
                })
            
            # Estadísticas de oraciones
            sentences = list(doc.sents)
            
            return {
                "entities": entities,
                "sentence_count": len(sentences),
                "avg_sentence_length": sum(len(sent.text.split()) for sent in sentences) / len(sentences) if sentences else 0,
                "language": doc.lang_,
                "entity_types": list(set(ent["label"] for ent in entities))
            }
            
        except Exception as e:
            return {"error": str(e)}
    
    async def _analyze_document_styles(self, file_path: str) -> Dict[str, Any]:
        """
        Análisis de estilos y formato del documento
        """
        try:
            doc = Document(file_path)
            
            # Analizar estilos utilizados
            styles_used = {}
            for para in doc.paragraphs:
                style_name = para.style.name if para.style else "Normal"
                styles_used[style_name] = styles_used.get(style_name, 0) + 1
            
            # Analizar fonts utilizados
            fonts_used = {}
            for para in doc.paragraphs:
                for run in para.runs:
                    if run.font.name:
                        fonts_used[run.font.name] = fonts_used.get(run.font.name, 0) + 1
            
            # Detectar problemas de formato
            formatting_issues = await self._detect_formatting_issues(doc)
            
            # Analizar consistencia
            consistency_score = await self._calculate_style_consistency(doc)
            
            return {
                "styles_used": styles_used,
                "fonts_used": fonts_used,
                "formatting_issues": formatting_issues,
                "consistency_score": consistency_score,
                "total_styles": len(styles_used),
                "total_fonts": len(fonts_used)
            }
            
        except Exception as e:
            return {"error": str(e)}
    
    async def _detect_formatting_issues(self, doc) -> List[str]:
        """
        Detectar problemas comunes de formato
        """
        issues = []
        
        # Detectar uso excesivo de espacios
        space_abuse_count = 0
        for para in doc.paragraphs:
            if "  " in para.text:  # Doble espacio
                space_abuse_count += 1
        
        if space_abuse_count > len(doc.paragraphs) * 0.1:  # Más del 10%
            issues.append("Excessive use of multiple spaces for alignment")
        
        # Detectar uso inconsistente de mayúsculas
        all_caps_count = 0
        for para in doc.paragraphs:
            if para.text.isupper() and len(para.text) > 10:
                all_caps_count += 1
        
        if all_caps_count > 0:
            issues.append(f"Found {all_caps_count} paragraphs in ALL CAPS")
        
        # Detectar párrafos muy largos
        long_paragraphs = sum(1 for para in doc.paragraphs if len(para.text) > 1000)
        if long_paragraphs > 0:
            issues.append(f"Found {long_paragraphs} very long paragraphs (>1000 chars)")
        
        return issues
    
    async def _calculate_style_consistency(self, doc) -> float:
        """
        Calcular puntuación de consistencia de estilos (0-1)
        """
        try:
            if not doc.paragraphs:
                return 1.0
            
            # Analizar consistencia de headings
            heading_styles = []
            for para in doc.paragraphs:
                style_name = para.style.name if para.style else "Normal"
                if "heading" in style_name.lower() or "title" in style_name.lower():
                    heading_styles.append(style_name)
            
            # Calcular varianza en estilos de heading
            heading_consistency = 1.0
            if heading_styles:
                unique_heading_styles = len(set(heading_styles))
                heading_consistency = max(0, 1 - (unique_heading_styles - 1) / len(heading_styles))
            
            # Analizar consistencia de fonts
            fonts_in_body = []
            for para in doc.paragraphs:
                for run in para.runs:
                    if run.font.name:
                        fonts_in_body.append(run.font.name)
            
            font_consistency = 1.0
            if fonts_in_body:
                unique_fonts = len(set(fonts_in_body))
                # Penalizar uso de más de 2-3 fonts diferentes
                if unique_fonts > 3:
                    font_consistency = max(0, 1 - (unique_fonts - 3) / len(fonts_in_body))
            
            # Promedio ponderado
            overall_consistency = (heading_consistency * 0.6 + font_consistency * 0.4)
            return round(overall_consistency, 3)
            
        except Exception:
            return 0.5  # Score neutral si hay error
    
    async def _analyze_embedded_tables(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Análisis especializado de tablas embebidas
        """
        try:
            # Usar docx2python para mejor extracción de tablas
            result = docx2python(file_path)
            
            tables_analysis = []
            
            # Analizar tablas extraídas
            if hasattr(result, 'body') and result.body:
                for section_idx, section in enumerate(result.body):
                    for table_idx, table_data in enumerate(section):
                        if isinstance(table_data, list) and len(table_data) > 1:
                            table_analysis = {
                                "section": section_idx,
                                "table_index": table_idx,
                                "rows": len(table_data),
                                "columns": len(table_data[0]) if table_data else 0,
                                "data_preview": table_data[:3],  # Primeras 3 filas
                                "has_header": self._detect_table_header(table_data),
                                "data_types": self._analyze_table_data_types(table_data)
                            }
                            tables_analysis.append(table_analysis)
            
            return tables_analysis
            
        except Exception as e:
            self.logger.error(f"Embedded tables analysis failed: {e}")
            return []
    
    def _detect_table_header(self, table_data: List[List]) -> bool:
        """
        Detectar si la tabla tiene fila de encabezado
        """
        if len(table_data) < 2:
            return False
        
        first_row = table_data[0]
        second_row = table_data[1]
        
        # Heurísticas para detectar header
        # 1. Primera fila no tiene números cuando segunda fila sí
        first_row_numeric = sum(1 for cell in first_row if self._is_numeric(str(cell)))
        second_row_numeric = sum(1 for cell in second_row if self._is_numeric(str(cell)))
        
        if first_row_numeric == 0 and second_row_numeric > 0:
            return True
        
        # 2. Primera fila es más corta en texto que el promedio
        if len(first_row) > 0:
            first_row_avg_length = sum(len(str(cell)) for cell in first_row) / len(first_row)
            other_rows_avg_length = 0
            
            if len(table_data) > 1:
                total_length = sum(len(str(cell)) for row in table_data[1:] for cell in row)
                total_cells = sum(len(row) for row in table_data[1:])
                other_rows_avg_length = total_length / total_cells if total_cells > 0 else 0
            
            return first_row_avg_length < other_rows_avg_length * 0.7
        
        return False
    
    def _analyze_table_data_types(self, table_data: List[List]) -> Dict[str, Any]:
        """
        Analizar tipos de datos en cada columna de la tabla
        """
        if not table_data:
            return {}
        
        max_cols = max(len(row) for row in table_data) if table_data else 0
        column_types = {}
        
        for col_idx in range(max_cols):
            column_values = []
            for row in table_data:
                if col_idx < len(row):
                    column_values.append(str(row[col_idx]).strip())
            
            # Analizar tipo predominante
            numeric_count = sum(1 for val in column_values if self._is_numeric(val) and val)
            date_count = sum(1 for val in column_values if self._is_date_like(val))
            email_count = sum(1 for val in column_values if self._is_email_like(val))
            
            total_non_empty = sum(1 for val in column_values if val)
            
            if total_non_empty == 0:
                column_type = "empty"
            elif numeric_count / total_non_empty > 0.7:
                column_type = "numeric"
            elif date_count / total_non_empty > 0.7:
                column_type = "date"
            elif email_count / total_non_empty > 0.7:
                column_type = "email"
            else:
                column_type = "text"
            
            column_types[f"column_{col_idx}"] = {
                "type": column_type,
                "confidence": max(numeric_count, date_count, email_count) / total_non_empty if total_non_empty > 0 else 0
            }
        
        return column_types
    
    def _is_date_like(self, text: str) -> bool:
        """Verificar si el texto parece una fecha"""
        date_patterns = [
            r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}',
            r'\d{4}[/-]\d{1,2}[/-]\d{1,2}',
            r'[A-Za-z]{3,9}\s+\d{1,2},?\s+\d{4}'
        ]
        
        return any(re.search(pattern, text) for pattern in date_patterns)
    
    def _is_email_like(self, text: str) -> bool:
        """Verificar si el texto parece un email"""
        email_pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
        return bool(re.match(email_pattern, text.strip()))
    
    async def _extract_document_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extraer metadatos del documento Word
        """
        try:
            doc = Document(file_path)
            core_props = doc.core_properties
            
            metadata = {
                "title": core_props.title,
                "author": core_props.author,
                "subject": core_props.subject,
                "keywords": core_props.keywords,
                "created": core_props.created.isoformat() if core_props.created else None,
                "modified": core_props.modified.isoformat() if core_props.modified else None,
                "last_modified_by": core_props.last_modified_by,
                "revision": core_props.revision,
                "version": core_props.version,
                "category": core_props.category,
                "comments": core_props.comments
            }
            
            # Estadísticas del documento
            stats = {
                "total_paragraphs": len(doc.paragraphs),
                "total_tables": len(doc.tables),
                "total_sections": len(doc.sections),
                "pages_estimate": self._estimate_page_count(doc)
            }
            
            metadata["document_stats"] = stats
            
            return metadata
            
        except Exception as e:
            return {"error": str(e)}
    
    def _estimate_page_count(self, doc) -> int:
        """
        Estimar número de páginas basándose en contenido
        """
        try:
            total_chars = sum(len(para.text) for para in doc.paragraphs)
            # Estimación: ~2500 caracteres por página (aproximado)
            estimated_pages = max(1, total_chars // 2500)
            return estimated_pages
        except:
            return 1
    
    async def _assess_document_quality(self, docx_analysis: Dict, content_analysis: Dict) -> Dict[str, Any]:
        """
        Evaluar calidad general del documento
        """
        try:
            quality_scores = {}
            
            # Puntuación de estructura (0-1)
            structure_score = self._calculate_structure_score(content_analysis)
            quality_scores["structure"] = structure_score
            
            # Puntuación de formato (0-1) 
            format_score = self._calculate_format_score(docx_analysis)
            quality_scores["formatting"] = format_score
            
            # Puntuación de contenido (0-1)
            content_score = self._calculate_content_score(content_analysis)
            quality_scores["content"] = content_score
            
            # Puntuación general
            overall_score = (structure_score + format_score + content_score) / 3
            quality_scores["overall"] = round(overall_score, 3)
            
            # Determinar nivel de calidad
            if overall_score >= 0.8:
                quality_level = "excellent"
            elif overall_score >= 0.6:
                quality_level = "good"
            elif overall_score >= 0.4:
                quality_level = "fair"
            else:
                quality_level = "poor"
            
            # Recomendaciones
            recommendations = self._generate_quality_recommendations(quality_scores, docx_analysis)
            
            return {
                "quality_scores": quality_scores,
                "quality_level": quality_level,
                "recommendations": recommendations
            }
            
        except Exception as e:
            return {
                "error": str(e),
                "quality_level": "unknown"
            }
    
    def _calculate_structure_score(self, content_analysis: Dict) -> float:
        """Calcular puntuación de estructura del documento"""
        score = 0.5  # Base score
        
        # Bonus por headings bien estructurados
        headings = content_analysis.get("headings", [])
        if headings:
            has_h1 = any(h["level"] == 1 for h in headings)
            has_hierarchy = len(set(h["level"] for h in headings)) > 1
            
            if has_h1:
                score += 0.2
            if has_hierarchy:
                score += 0.2
        
        # Bonus por listas organizadas
        lists = content_analysis.get("lists", [])
        if lists:
            score += min(len(lists) * 0.05, 0.1)
        
        return min(score, 1.0)
    
    def _calculate_format_score(self, docx_analysis: Dict) -> float:
        """Calcular puntuación de formato del documento"""
        score = 0.5  # Base score
        
        # Penalizar por problemas de formato
        formatting_issues = docx_analysis.get("styles", {}).get("formatting_issues", [])
        if formatting_issues:
            penalty = len(formatting_issues) * 0.1
            score -= penalty
        
        # Bonus por consistencia de estilos
        consistency_score = docx_analysis.get("styles", {}).get("consistency_score", 0.5)
        score = (score + consistency_score) / 2
        
        return max(0.0, min(score, 1.0))
    
    def _calculate_content_score(self, content_analysis: Dict) -> float:
        """Calcular puntuación de calidad del contenido"""
        score = 0.5  # Base score
        
        word_count = content_analysis.get("word_count", 0)
        
        # Bonus por longitud apropiada
        if 100 <= word_count <= 10000:  # Rango razonable
            score += 0.2
        elif word_count > 50:  # Al menos algo de contenido
            score += 0.1
        
        # Bonus por análisis NLP exitoso
        nlp_analysis = content_analysis.get("nlp_analysis", {})
        if nlp_analysis and "entities" in nlp_analysis:
            entities = nlp_analysis["entities"]
            if entities:
                score += min(len(entities) * 0.02, 0.2)
        
        return min(score, 1.0)
    
    def _generate_quality_recommendations(self, quality_scores: Dict, docx_analysis: Dict) -> List[str]:
        """Generar recomendaciones para mejorar calidad del documento"""
        recommendations = []
        
        # Recomendaciones de estructura
        if quality_scores.get("structure", 0) < 0.6:
            recommendations.append("Consider adding more headings and structure to improve document organization")
        
        # Recomendaciones de formato
        if quality_scores.get("formatting", 0) < 0.6:
            recommendations.append("Review and standardize document formatting for better consistency")
            
            formatting_issues = docx_analysis.get("styles", {}).get("formatting_issues", [])
            if formatting_issues:
                recommendations.append(f"Address {len(formatting_issues)} formatting issues identified")
        
        # Recomendaciones de contenido
        if quality_scores.get("content", 0) < 0.6:
            recommendations.append("Consider expanding content or improving text quality")
        
        if not recommendations:
            recommendations.append("Document quality is good - no major improvements needed")
        
        return recommendations
    
    def _get_alignment_name(self, alignment) -> str:
        """Convertir alignment enum a string"""
        try:
            if alignment == WD_ALIGN_PARAGRAPH.LEFT:
                return "left"
            elif alignment == WD_ALIGN_PARAGRAPH.CENTER:
                return "center"
            elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                return "right"
            elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                return "justify"
            else:
                return "left"  # default
        except:
            return "left"
    
    def _get_analysis_capabilities(self) -> Dict[str, bool]:
        """Obtener capacidades de análisis disponibles"""
        return {
            "python_docx": True,
            "docx2txt": True,
            "mammoth": True,
            "docx2python": True,
            "spacy_nlp": self.nlp is not None,
            "html_parsing": True,
            "style_analysis": True,
            "table_analysis": True,
            "metadata_extraction": True,
            "quality_assessment": True
        }
    
    async def extract_text_for_conversion(self, file_path: str) -> Dict[str, Any]:
        """
        Extraer texto estructurado para conversión a otros formatos
        """
        try:
            analysis = await self.analyze_word_document(file_path)
            
            # Extraer elementos estructurados
            structured_content = {
                "paragraphs": [],
                "tables": [],
                "headings": [],
                "lists": []
            }
            
            # Procesar paragraphs
            docx_analysis = analysis.get("analysis_results", {}).get("structure", {})
            for para in docx_analysis.get("paragraphs", []):
                if para["text"].strip():
                    structured_content["paragraphs"].append({
                        "text": para["text"],
                        "style": para["style"],
                        "formatting": {
                            "alignment": para.get("alignment", "left"),
                            "runs": para.get("runs", [])
                        }
                    })
            
            # Procesar tablas
            for table in docx_analysis.get("tables", []):
                structured_content["tables"].append({
                    "index": table["table_index"],
                    "rows": table["rows"],
                    "columns": table["columns"],
                    "data": table["data"],
                    "type": table.get("table_type", "generic")
                })
            
            # Procesar headings
            content_analysis = analysis.get("analysis_results", {}).get("content", {})
            structured_content["headings"] = content_analysis.get("headings", [])
            structured_content["lists"] = content_analysis.get("lists", [])
            
            return {
                "status": "success",
                "structured_content": structured_content,
                "metadata": analysis.get("analysis_results", {}).get("metadata", {}),
                "quality_info": analysis.get("analysis_results", {}).get("quality_assessment", {})
            }
            
        except Exception as e:
            return {
                "status": "failed",
                "error": str(e)
            }