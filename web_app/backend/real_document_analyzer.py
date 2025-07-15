#!/usr/bin/env python3
"""
Real Document Intelligence Analyzer
Análisis real de documentos con procesamiento de contenido
"""

import fitz  # PyMuPDF
from docx import Document
import openpyxl
from PIL import Image
import io
import re
from typing import Dict, List, Any, Tuple
import tempfile
import os
from pathlib import Path
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class RealDocumentAnalyzer:
    """Real document analyzer with actual content processing"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self.analyze_pdf,
            '.docx': self.analyze_word,
            '.doc': self.analyze_word,
            '.xlsx': self.analyze_excel,
            '.xls': self.analyze_excel
        }
    
    def analyze_document(self, file_path: str, filename: str) -> Dict[str, Any]:
        """Analyze a document and extract real content"""
        file_ext = Path(filename).suffix.lower()
        file_size = os.path.getsize(file_path)
        
        if file_ext not in self.supported_formats:
            return self._create_unsupported_result(filename, file_size, file_ext)
        
        try:
            analyzer_func = self.supported_formats[file_ext]
            analysis_result = analyzer_func(file_path)
            
            # Add common metadata
            analysis_result.update({
                "filename": filename,
                "file_size": file_size,
                "file_format": file_ext,
                "status": "completed"
            })
            
            return analysis_result
            
        except Exception as e:
            logger.error(f"Error analyzing {filename}: {str(e)}")
            return self._create_error_result(filename, file_size, str(e))
    
    def analyze_pdf(self, file_path: str) -> Dict[str, Any]:
        """Real PDF analysis with text, images, and table detection"""
        doc = fitz.open(file_path)
        
        # Basic document info
        page_count = len(doc)
        text_content = ""
        images = []
        tables_detected = 0
        text_blocks = []
        
        # Process each page
        for page_num in range(page_count):
            page = doc[page_num]
            
            # Extract text
            page_text = page.get_text()
            text_content += page_text + "\n"
            
            # Extract text blocks with positions
            blocks = page.get_text("dict")
            for block in blocks.get("blocks", []):
                if "lines" in block:
                    for line in block["lines"]:
                        for span in line.get("spans", []):
                            text_blocks.append({
                                "text": span.get("text", ""),
                                "font": span.get("font", ""),
                                "size": span.get("size", 0),
                                "bbox": span.get("bbox", [])
                            })
            
            # Detect tables (simple heuristic)
            tables_detected += self._detect_tables_in_text(page_text)
            
            # Extract images
            image_list = page.get_images()
            images.extend([{
                "page": page_num + 1,
                "index": img[0],
                "width": img[2] if len(img) > 2 else 0,
                "height": img[3] if len(img) > 3 else 0
            } for img in image_list])
        
        doc.close()
        
        # Analyze content quality
        word_count = len(text_content.split())
        char_count = len(text_content)
        
        # Detect document type based on content
        doc_type = self._classify_pdf_content(text_content)
        
        # Extract key information
        extracted_data = self._extract_key_data(text_content)
        
        # Calculate confidence based on text quality
        confidence = self._calculate_pdf_confidence(text_content, word_count, page_count)
        
        return {
            "document_type": "PDF",
            "confidence_score": confidence,
            "analysis": {
                "page_count": page_count,
                "word_count": word_count,
                "character_count": char_count,
                "images_found": len(images),
                "tables_detected": tables_detected,
                "text_blocks": len(text_blocks),
                "estimated_processing_time": f"{page_count * 0.5:.1f} seconds",
                "document_classification": doc_type,
                "content_preview": text_content[:500] + "..." if len(text_content) > 500 else text_content,
                "extracted_data": extracted_data,
                "text_quality": self._assess_text_quality(text_content),
                "images": images[:5],  # First 5 images only
                "structure_analysis": {
                    "has_headers": self._has_headers(text_blocks),
                    "has_tables": tables_detected > 0,
                    "has_images": len(images) > 0,
                    "language_detected": self._detect_language(text_content)
                }
            }
        }
    
    def analyze_word(self, file_path: str) -> Dict[str, Any]:
        """Real Word document analysis"""
        try:
            doc = Document(file_path)
            
            # Extract text content
            full_text = []
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            text_content = '\n'.join(full_text)
            
            # Analyze structure
            paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
            
            # Count tables
            table_count = len(doc.tables)
            table_data = []
            for table in doc.tables[:3]:  # Analyze first 3 tables
                rows = len(table.rows)
                cols = len(table.columns) if table.rows else 0
                table_data.append({"rows": rows, "columns": cols})
            
            # Analyze styles and formatting
            styles_used = set()
            for paragraph in doc.paragraphs:
                styles_used.add(paragraph.style.name)
            
            # Word count and statistics
            word_count = len(text_content.split())
            char_count = len(text_content)
            
            # Extract headers (assuming Heading styles)
            headers = []
            for paragraph in doc.paragraphs:
                if 'Heading' in paragraph.style.name and paragraph.text.strip():
                    headers.append({
                        "text": paragraph.text,
                        "level": paragraph.style.name
                    })
            
            # Extract key data
            extracted_data = self._extract_key_data(text_content)
            
            # Document classification
            doc_type = self._classify_document_content(text_content)
            
            confidence = self._calculate_word_confidence(text_content, word_count, paragraph_count)
            
            return {
                "document_type": "Word Document",
                "confidence_score": confidence,
                "analysis": {
                    "word_count": word_count,
                    "character_count": char_count,
                    "paragraph_count": paragraph_count,
                    "table_count": table_count,
                    "styles_count": len(styles_used),
                    "headers_count": len(headers),
                    "estimated_processing_time": f"{word_count / 1000:.1f} seconds",
                    "document_classification": doc_type,
                    "content_preview": text_content[:500] + "..." if len(text_content) > 500 else text_content,
                    "extracted_data": extracted_data,
                    "structure_analysis": {
                        "headers": headers[:5],  # First 5 headers
                        "tables": table_data,
                        "styles_used": list(styles_used)[:10],  # First 10 styles
                        "has_structured_content": len(headers) > 0 or table_count > 0,
                        "language_detected": self._detect_language(text_content)
                    },
                    "quality_assessment": {
                        "text_quality": self._assess_text_quality(text_content),
                        "structure_quality": "good" if len(headers) > 0 else "basic",
                        "formatting_complexity": "high" if len(styles_used) > 5 else "basic"
                    }
                }
            }
            
        except Exception as e:
            logger.error(f"Error analyzing Word document: {str(e)}")
            return self._create_error_result("word_document", 0, str(e))
    
    def analyze_excel(self, file_path: str) -> Dict[str, Any]:
        """Real Excel file analysis"""
        try:
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            
            sheet_names = workbook.sheetnames
            sheet_count = len(sheet_names)
            
            total_rows = 0
            total_cols = 0
            total_cells_with_data = 0
            formulas_found = 0
            charts_found = 0
            
            sheets_analysis = []
            
            # Analyze each sheet
            for sheet_name in sheet_names[:5]:  # Analyze first 5 sheets
                sheet = workbook[sheet_name]
                
                # Get used range
                max_row = sheet.max_row
                max_col = sheet.max_column
                
                total_rows += max_row
                total_cols += max_col
                
                # Count cells with data and formulas
                cells_with_data = 0
                sheet_formulas = 0
                sample_data = []
                
                for row in range(1, min(max_row + 1, 50)):  # Sample first 50 rows
                    for col in range(1, min(max_col + 1, 20)):  # Sample first 20 columns
                        cell = sheet.cell(row=row, column=col)
                        if cell.value is not None:
                            cells_with_data += 1
                            if row <= 5 and col <= 5:  # Sample data for preview
                                sample_data.append({
                                    "row": row,
                                    "col": col,
                                    "value": str(cell.value)[:50]  # Truncate long values
                                })
                            
                            # Check for formulas
                            if str(cell.value).startswith('='):
                                sheet_formulas += 1
                
                total_cells_with_data += cells_with_data
                formulas_found += sheet_formulas
                
                # Detect data patterns
                data_patterns = self._detect_excel_patterns(sheet, max_row, max_col)
                
                sheets_analysis.append({
                    "name": sheet_name,
                    "rows": max_row,
                    "columns": max_col,
                    "cells_with_data": cells_with_data,
                    "formulas": sheet_formulas,
                    "sample_data": sample_data[:10],  # First 10 cells
                    "data_patterns": data_patterns
                })
            
            workbook.close()
            
            # Calculate confidence
            confidence = self._calculate_excel_confidence(total_cells_with_data, formulas_found, sheet_count)
            
            # Classify the Excel file type
            file_classification = self._classify_excel_content(sheets_analysis)
            
            return {
                "document_type": "Excel Spreadsheet",
                "confidence_score": confidence,
                "analysis": {
                    "sheet_count": sheet_count,
                    "total_rows": total_rows,
                    "total_columns": total_cols,
                    "cells_with_data": total_cells_with_data,
                    "formulas_found": formulas_found,
                    "charts_found": charts_found,
                    "estimated_processing_time": f"{sheet_count * 0.8:.1f} seconds",
                    "file_classification": file_classification,
                    "sheets_analysis": sheets_analysis,
                    "data_summary": {
                        "avg_rows_per_sheet": total_rows / sheet_count if sheet_count > 0 else 0,
                        "avg_cols_per_sheet": total_cols / sheet_count if sheet_count > 0 else 0,
                        "data_density": total_cells_with_data / (total_rows * total_cols) if total_rows * total_cols > 0 else 0,
                        "has_formulas": formulas_found > 0,
                        "complexity": "high" if formulas_found > 10 else "medium" if formulas_found > 0 else "basic"
                    }
                }
            }
            
        except Exception as e:
            logger.error(f"Error analyzing Excel file: {str(e)}")
            return self._create_error_result("excel_file", 0, str(e))
    
    # Helper methods for analysis
    def _detect_tables_in_text(self, text: str) -> int:
        """Simple table detection in text"""
        lines = text.split('\n')
        table_indicators = 0
        
        for line in lines:
            # Look for common table patterns
            if '\t' in line or '|' in line:
                table_indicators += 1
            elif re.search(r'\s{3,}', line) and len(line.split()) > 2:
                table_indicators += 1
        
        return min(table_indicators // 3, 5)  # Estimate number of tables
    
    def _classify_pdf_content(self, text: str) -> str:
        """Classify PDF based on content"""
        text_lower = text.lower()
        
        if any(word in text_lower for word in ['invoice', 'bill', 'amount', 'total', 'payment']):
            return "invoice/bill"
        elif any(word in text_lower for word in ['contract', 'agreement', 'terms', 'conditions']):
            return "contract/legal"
        elif any(word in text_lower for word in ['report', 'analysis', 'summary', 'findings']):
            return "report/analysis"
        elif any(word in text_lower for word in ['manual', 'instructions', 'guide', 'how to']):
            return "manual/guide"
        elif any(word in text_lower for word in ['resume', 'cv', 'experience', 'education']):
            return "resume/cv"
        else:
            return "general document"
    
    def _classify_document_content(self, text: str) -> str:
        """Classify document based on content"""
        return self._classify_pdf_content(text)  # Use same logic
    
    def _classify_excel_content(self, sheets_analysis: List[Dict]) -> str:
        """Classify Excel file based on content"""
        total_formulas = sum(sheet.get('formulas', 0) for sheet in sheets_analysis)
        
        if total_formulas > 20:
            return "financial/calculation model"
        elif any('budget' in sheet['name'].lower() for sheet in sheets_analysis):
            return "budget/planning"
        elif any('data' in sheet['name'].lower() for sheet in sheets_analysis):
            return "data analysis"
        elif len(sheets_analysis) > 1:
            return "multi-sheet workbook"
        else:
            return "data spreadsheet"
    
    def _extract_key_data(self, text: str) -> Dict[str, List[str]]:
        """Extract key data like emails, dates, phones, etc."""
        
        # Email pattern
        emails = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
        
        # Date patterns (various formats)
        dates = re.findall(r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b', text)
        
        # Phone patterns
        phones = re.findall(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b|\(\d{3}\)\s*\d{3}[-.]?\d{4}\b', text)
        
        # Currency amounts
        amounts = re.findall(r'\$\d{1,3}(?:,\d{3})*(?:\.\d{2})?|\b\d{1,3}(?:,\d{3})*(?:\.\d{2})?\s*(?:USD|EUR|GBP)\b', text)
        
        return {
            "emails": list(set(emails))[:5],
            "dates": list(set(dates))[:5],
            "phones": list(set(phones))[:5],
            "amounts": list(set(amounts))[:5]
        }
    
    def _detect_language(self, text: str) -> str:
        """Simple language detection"""
        # Basic language detection based on common words
        spanish_words = ['el', 'la', 'de', 'que', 'y', 'en', 'un', 'es', 'se', 'no', 'te', 'lo', 'le', 'da', 'su', 'por', 'son', 'con', 'para', 'al']
        english_words = ['the', 'be', 'to', 'of', 'and', 'a', 'in', 'that', 'have', 'it', 'for', 'not', 'on', 'with', 'as', 'you', 'do', 'at', 'this', 'but']
        
        text_words = text.lower().split()
        spanish_count = sum(1 for word in text_words if word in spanish_words)
        english_count = sum(1 for word in text_words if word in english_words)
        
        if spanish_count > english_count:
            return "spanish"
        elif english_count > spanish_count:
            return "english"
        else:
            return "unknown"
    
    def _assess_text_quality(self, text: str) -> str:
        """Assess the quality of extracted text"""
        if len(text) < 50:
            return "poor"
        elif len(text.split()) / len(text) > 0.2:  # Good word density
            return "excellent"
        elif len(text.split()) / len(text) > 0.1:
            return "good"
        else:
            return "fair"
    
    def _has_headers(self, text_blocks: List[Dict]) -> bool:
        """Check if document has clear headers"""
        if not text_blocks:
            return False
        
        # Look for larger font sizes that might indicate headers
        font_sizes = [block.get('size', 0) for block in text_blocks if block.get('size', 0) > 0]
        if not font_sizes:
            return False
        
        avg_size = sum(font_sizes) / len(font_sizes)
        return any(size > avg_size * 1.2 for size in font_sizes)
    
    def _detect_excel_patterns(self, sheet, max_row: int, max_col: int) -> Dict[str, Any]:
        """Detect data patterns in Excel sheet"""
        patterns = {
            "has_headers": False,
            "numeric_columns": 0,
            "text_columns": 0,
            "date_columns": 0,
            "empty_columns": 0
        }
        
        # Check first row for headers
        first_row_text = 0
        for col in range(1, min(max_col + 1, 10)):
            cell_value = sheet.cell(row=1, column=col).value
            if isinstance(cell_value, str) and cell_value.strip():
                first_row_text += 1
        
        patterns["has_headers"] = first_row_text > max_col * 0.5
        
        # Analyze column types (sample first 10 columns, 20 rows)
        for col in range(1, min(max_col + 1, 10)):
            numeric_count = 0
            text_count = 0
            date_count = 0
            empty_count = 0
            
            for row in range(2, min(max_row + 1, 20)):  # Skip header row
                cell_value = sheet.cell(row=row, column=col).value
                
                if cell_value is None:
                    empty_count += 1
                elif isinstance(cell_value, (int, float)):
                    numeric_count += 1
                elif isinstance(cell_value, str):
                    text_count += 1
                else:  # datetime or other
                    date_count += 1
            
            # Classify column based on majority type
            total_checked = min(18, max_row - 1)
            if total_checked > 0:
                if numeric_count / total_checked > 0.7:
                    patterns["numeric_columns"] += 1
                elif text_count / total_checked > 0.7:
                    patterns["text_columns"] += 1
                elif date_count / total_checked > 0.7:
                    patterns["date_columns"] += 1
                elif empty_count / total_checked > 0.7:
                    patterns["empty_columns"] += 1
        
        return patterns
    
    def _calculate_pdf_confidence(self, text: str, word_count: int, page_count: int) -> float:
        """Calculate confidence score for PDF analysis"""
        confidence = 0.5  # Base confidence
        
        # Text quality factors
        if word_count > 100:
            confidence += 0.2
        if len(text) > 500:
            confidence += 0.1
        if word_count / page_count > 50:  # Good text density per page
            confidence += 0.1
        
        # Content structure factors
        if self._extract_key_data(text)['emails']:
            confidence += 0.05
        if self._extract_key_data(text)['dates']:
            confidence += 0.05
        
        return min(confidence, 0.95)
    
    def _calculate_word_confidence(self, text: str, word_count: int, paragraph_count: int) -> float:
        """Calculate confidence score for Word analysis"""
        confidence = 0.7  # Higher base for Word docs
        
        if word_count > 200:
            confidence += 0.1
        if paragraph_count > 5:
            confidence += 0.1
        if word_count / paragraph_count > 10:  # Good paragraph density
            confidence += 0.05
        
        return min(confidence, 0.95)
    
    def _calculate_excel_confidence(self, cells_with_data: int, formulas: int, sheets: int) -> float:
        """Calculate confidence score for Excel analysis"""
        confidence = 0.6  # Base confidence
        
        if cells_with_data > 50:
            confidence += 0.15
        if formulas > 0:
            confidence += 0.1
        if sheets > 1:
            confidence += 0.05
        if cells_with_data > 200:
            confidence += 0.1
        
        return min(confidence, 0.95)
    
    def _create_unsupported_result(self, filename: str, file_size: int, file_ext: str) -> Dict[str, Any]:
        """Create result for unsupported file format"""
        return {
            "filename": filename,
            "document_type": "Unsupported Format",
            "file_size": file_size,
            "confidence_score": 0,
            "status": "unsupported",
            "analysis": {
                "error": f"File format {file_ext} not supported",
                "supported_formats": [".pdf", ".docx", ".doc", ".xlsx", ".xls"]
            }
        }
    
    def _create_error_result(self, filename: str, file_size: int, error_msg: str) -> Dict[str, Any]:
        """Create result for analysis errors"""
        return {
            "filename": filename,
            "document_type": "Error",
            "file_size": file_size,
            "confidence_score": 0,
            "status": "error",
            "analysis": {
                "error": error_msg,
                "estimated_processing_time": "N/A"
            }
        }