#!/usr/bin/env python3
"""
Create test documents for real analysis
"""

from docx import Document
import openpyxl
from openpyxl import Workbook

def create_test_word_doc():
    """Create a sample Word document"""
    doc = Document()
    
    # Add title
    title = doc.add_heading('Sample Business Report', 0)
    
    # Add introduction paragraph
    intro = doc.add_paragraph(
        'This is a comprehensive business report analyzing market trends and financial performance. '
        'The document contains structured information including tables, headings, and contact details.'
    )
    
    # Add contact information
    doc.add_heading('Contact Information', level=1)
    contact_para = doc.add_paragraph()
    contact_para.add_run('Company: ').bold = True
    contact_para.add_run('TechCorp Solutions Inc.\n')
    contact_para.add_run('Email: ').bold = True
    contact_para.add_run('contact@techcorp.com\n')
    contact_para.add_run('Phone: ').bold = True
    contact_para.add_run('+1-555-987-6543\n')
    contact_para.add_run('Date: ').bold = True
    contact_para.add_run('December 15, 2024')
    
    # Add financial section
    doc.add_heading('Financial Overview', level=1)
    doc.add_paragraph(
        'Our quarterly revenue analysis shows significant growth in key market segments. '
        'The total revenue for Q4 2024 reached $2,500,000 with a profit margin of 15.5%.'
    )
    
    # Add a table
    doc.add_heading('Quarterly Performance', level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # Add header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Quarter'
    hdr_cells[1].text = 'Revenue'
    hdr_cells[2].text = 'Expenses'
    hdr_cells[3].text = 'Profit'
    
    # Add data rows
    quarters = [
        ('Q1 2024', '$1,800,000', '$1,400,000', '$400,000'),
        ('Q2 2024', '$2,100,000', '$1,600,000', '$500,000'),
        ('Q3 2024', '$2,300,000', '$1,750,000', '$550,000'),
        ('Q4 2024', '$2,500,000', '$1,900,000', '$600,000')
    ]
    
    for quarter, revenue, expenses, profit in quarters:
        row_cells = table.add_row().cells
        row_cells[0].text = quarter
        row_cells[1].text = revenue
        row_cells[2].text = expenses
        row_cells[3].text = profit
    
    # Add conclusion
    doc.add_heading('Conclusions', level=1)
    doc.add_paragraph(
        'The analysis demonstrates consistent growth throughout 2024. '
        'Key recommendations include expanding into emerging markets and '
        'investing in technology infrastructure for sustained growth.'
    )
    
    # Save document
    doc.save('/tmp/sample_business_report.docx')
    print("✅ Created sample Word document: /tmp/sample_business_report.docx")

def create_test_excel_file():
    """Create a sample Excel file"""
    wb = Workbook()
    ws = wb.active
    ws.title = "Sales Data"
    
    # Add headers
    headers = ['Product', 'Q1 Sales', 'Q2 Sales', 'Q3 Sales', 'Q4 Sales', 'Total', 'Growth %']
    for col, header in enumerate(headers, 1):
        ws.cell(row=1, column=col, value=header)
    
    # Add sample data
    products = [
        ['Product A', 15000, 18000, 22000, 25000],
        ['Product B', 12000, 14000, 16000, 18000],
        ['Product C', 8000, 9500, 11000, 13000],
        ['Product D', 20000, 22000, 24000, 27000],
        ['Product E', 5000, 6000, 7500, 9000]
    ]
    
    for row, (product, q1, q2, q3, q4) in enumerate(products, 2):
        ws.cell(row=row, column=1, value=product)
        ws.cell(row=row, column=2, value=q1)
        ws.cell(row=row, column=3, value=q2)
        ws.cell(row=row, column=4, value=q3)
        ws.cell(row=row, column=5, value=q4)
        
        # Add total formula
        ws.cell(row=row, column=6, value=f'=B{row}+C{row}+D{row}+E{row}')
        
        # Add growth formula
        ws.cell(row=row, column=7, value=f'=(E{row}-B{row})/B{row}*100')
    
    # Add summary row
    summary_row = len(products) + 3
    ws.cell(row=summary_row, column=1, value='TOTAL')
    for col in range(2, 7):
        ws.cell(row=summary_row, column=col, value=f'=SUM({chr(65+col-1)}2:{chr(65+col-1)}{len(products)+1})')
    
    # Create second sheet for customer data
    ws2 = wb.create_sheet("Customer Data")
    customer_headers = ['Customer ID', 'Name', 'Email', 'Phone', 'Total Orders', 'Revenue']
    for col, header in enumerate(customer_headers, 1):
        ws2.cell(row=1, column=col, value=header)
    
    customers = [
        ['CUST001', 'ABC Corporation', 'orders@abc-corp.com', '555-0101', 15, 45000],
        ['CUST002', 'XYZ Ltd', 'purchasing@xyz-ltd.com', '555-0102', 22, 67000],
        ['CUST003', 'Tech Innovations', 'sales@techinno.com', '555-0103', 8, 23000],
        ['CUST004', 'Global Systems', 'orders@globalsys.com', '555-0104', 31, 89000]
    ]
    
    for row, (cust_id, name, email, phone, orders, revenue) in enumerate(customers, 2):
        ws2.cell(row=row, column=1, value=cust_id)
        ws2.cell(row=row, column=2, value=name)
        ws2.cell(row=row, column=3, value=email)
        ws2.cell(row=row, column=4, value=phone)
        ws2.cell(row=row, column=5, value=orders)
        ws2.cell(row=row, column=6, value=revenue)
    
    # Save workbook
    wb.save('/tmp/sample_sales_data.xlsx')
    print("✅ Created sample Excel file: /tmp/sample_sales_data.xlsx")

if __name__ == "__main__":
    create_test_word_doc()
    create_test_excel_file()
    print("\n🎉 Test documents created successfully!")
    print("📄 Word: /tmp/sample_business_report.docx")
    print("📊 Excel: /tmp/sample_sales_data.xlsx")